# -*- coding: utf-8 -*-
# routers/assessment_router.py
# Production Final Version - 2026 Optimized for DB Persistence & Professional Reporting

import os
import re
import uuid
import json
import asyncio
import logging
import mimetypes
import tempfile
import pytz
from datetime import datetime
from typing import Optional, Dict, Any, Union, List

from fastapi import APIRouter, BackgroundTasks, HTTPException, Depends, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel

# --- Docx Imports (สำหรับสร้าง Report) ---
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# --- Project Imports ---
from routers.auth_router import get_current_user, check_user_permission, UserMe
from utils.path_utils import (
    _n, 
    get_tenant_year_export_root, 
    load_doc_id_mapping, 
    get_document_file_path,
    get_vectorstore_collection_path,
    get_vectorstore_tenant_root_path
)

from core.seam_assessment import SEAMPDCAEngine, AssessmentConfig
from core.vectorstore import load_all_vectorstores
from models.llm import create_llm_instance
from config.global_vars import (
    EVIDENCE_DOC_TYPES, 
    DEFAULT_LLM_MODEL_NAME, 
    DEFAULT_YEAR, 
    DEFAULT_TENANT,
    DATA_STORE_ROOT
)

# 🎯 Database Components (SQLite Persistence)
from database import (
    SessionLocal, 
    AssessmentTaskTable, 
    AssessmentResultTable,
    db_update_task_status,
    db_finish_task
)

MAX_CONCURRENT_TASKS = 4 
assessment_semaphore = asyncio.Semaphore(MAX_CONCURRENT_TASKS)


# ตั้งค่า Logger และ Router
logger = logging.getLogger(__name__)
assessment_router = APIRouter(prefix="/api/assess", tags=["Assessment"])

# --- Request Models ---
class StartAssessmentRequest(BaseModel):
    tenant: str
    year: Optional[Union[int, str]] = None
    enabler: str = "KM"
    sub_criteria: Optional[str] = "all"
    sequential_mode: bool = False

# ------------------------------------------------------------------
# [Helpers]
# ------------------------------------------------------------------
def parse_safe_date(raw_date_str: Any, file_path: str) -> str:
    """แปลงวันที่จาก String หรือ File Metadata ให้เป็น ISO Format (Bangkok Time)"""
    tz = pytz.timezone('Asia/Bangkok')
    if raw_date_str and isinstance(raw_date_str, str):
        try:
            # รองรับ format yyyymmdd_hhmmss
            if "_" in raw_date_str:
                dt = datetime.strptime(raw_date_str, "%Y%m%d_%H%M%S")
                return tz.localize(dt).isoformat()
        except: pass
    
    # Fallback: ใช้เวลาแก้ไขไฟล์ล่าสุด
    try:
        mtime = os.path.getmtime(file_path)
        dt = datetime.fromtimestamp(mtime, tz)
        return dt.isoformat()
    except:
        return datetime.now(tz).isoformat()


def safe_float(value):
    try:
        if isinstance(value, str):
            value = value.replace('%', '') # ลบ % ออกถ้ามี
        return float(value)
    except:
        return 0.0
    

@assessment_router.get("/evidence/{doc_type}/{document_uuid}")
async def serve_evidence_file(
    document_uuid: str,
    doc_type: str,
    tenant: str,
    year: str = None,
    enabler: str = None,
    current_user: UserMe = Depends(get_current_user)
):
    check_user_permission(current_user, tenant, enabler or "KM")

    file_info = get_document_file_path(
        document_uuid=document_uuid,
        tenant=tenant,
        year=year,
        enabler=enabler,
        doc_type_name=doc_type
    )

    if not file_info:
        raise HTTPException(status_code=404, detail="File not found")

    file_path = file_info["file_path"]
    
    # ดึงนามสกุลไฟล์
    ext = os.path.splitext(file_path)[1].lower()
    
    # 🛡️ Force MIME Type สำหรับ Mac/Safari
    mime_map = {
        ".pdf": "application/pdf",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
    }
    
    mime_type = mime_map.get(ext) or mimetypes.guess_type(file_path)[0] or "application/octet-stream"

    # ส่ง FileResponse
    response = FileResponse(
        path=file_path,
        media_type=mime_type,
        content_disposition_type="inline"
    )

    # 💡 หัวใจสำคัญสำหรับ Mac/Safari:
    # 1. ป้องกันไม่ให้ Browser ใช้ชื่อไฟล์จาก Path ซึ่งบางทีมีภาษาไทยแล้วทำให้ Header เพี้ยน
    # 2. บังคับ Header ให้ชัดเจน
    response.headers["Content-Type"] = mime_type
    response.headers["Accept-Ranges"] = "bytes" 
    
    # ถ้าเป็น PDF บน Mac ให้เติม Cache-Control เพื่อให้ Viewer ทำงานได้ดีขึ้น
    if ext == ".pdf":
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"

    return response


@assessment_router.get("/view-document")
async def view_document(
    filename: str, 
    document_uuid: Optional[str] = None, 
    current_user: UserMe = Depends(get_current_user)
):
    """ Endpoint สำหรับเปิดไฟล์ PDF โดยรองรับการค้นหาที่แม่นยำขึ้น """
    
    file_path = None

    # 1. พยายามหาจาก UUID ก่อน (ดีที่สุด เพราะแม่นยำ 100% แม้ชื่อไฟล์จะซ้ำ)
    if document_uuid:
        file_info = get_document_file_path(
            document_uuid=document_uuid,
            tenant=current_user.tenant,
            year=current_user.year,
            enabler="KM", # หรือดึงจาก Query Param
            doc_type_name="evidence"
        )
        if file_info:
            file_path = file_info["file_path"]

    # 2. ถ้าไม่มี UUID ให้หาจากชื่อไฟล์ (ใช้ Logic เดิมของคุณ)
    if not file_path:
        # ใช้ get_document_source_dir ที่คุณมีใน path_utils อยู่แล้ว
        from utils.path_utils import get_document_source_dir, resolve_filepath_to_absolute
        
        base_path = get_document_source_dir(
            tenant=current_user.tenant,
            year=current_user.year,
            enabler="KM",
            doc_type="evidence"
        )
        file_path = resolve_filepath_to_absolute(os.path.join(base_path, filename))

    if not file_path or not os.path.exists(file_path):
        logger.error(f"❌ File not found: {file_path}")
        raise HTTPException(status_code=404, detail=f"ไม่พบไฟล์เอกสารบนเซิร์ฟเวอร์")

    # ส่งไฟล์ PDF กลับไป
    return FileResponse(file_path, media_type="application/pdf")


def _transform_result_for_ui(raw_data: Dict[str, Any], current_user: Any = None) -> Dict[str, Any]:
    if not raw_data or not isinstance(raw_data, dict):
        return {"status": "FAILED", "message": "Invalid data format"}

    # --- [0] RESOLVE CORE DATA ---
    res = raw_data.get("result") or raw_data.get("assessment_result") or raw_data
    metadata = res.get("metadata", {})
    summary = res.get("summary") or res.get("result_summary", {})
    global_evidence_map = raw_data.get("evidence_map") or res.get("evidence_map") or {}

    # --- [1] ENABLER ROADMAP ---
    raw_global_roadmap = res.get("enabler_roadmap") or res.get("strategic_roadmap") or {}
    ui_global_roadmap = {
        "status": raw_global_roadmap.get("status", "SUCCESS"),
        "overall_strategy": raw_global_roadmap.get("overall_strategy") or "มุ่งเน้นการยกระดับตามมาตรฐาน SE-AM",
        "phases": raw_global_roadmap.get("phases") or []
    }

    # --- [2] SUB-CRITERIA PROCESSING ---
    processed = []
    radar_data = []
    sub_list = res.get("sub_criteria_details") or res.get("sub_criteria_results") or []

    for sub in sub_list:
        sub_id = str(sub.get("sub_id", "N/A"))
        lv_details = sub.get("level_details", {}) or {}
        
        sub_roadmap_data = sub.get("sub_roadmap") or {}
        ui_sub_roadmap = {
            "strategy": sub_roadmap_data.get("overall_strategy") or sub.get("strategic_focus", ""),
            "phases": sub_roadmap_data.get("phases") or [],
            "is_gap_detected": sub_roadmap_data.get("is_gap_detected", False)
        }

        ui_levels = {}
        pdca_matrix = []
        pdca_coverage = {}
        grouped_sources = {str(i): [] for i in range(1, 6)}
        
        # 🚩 FIX ISSUE 2: ย้าย pool มาไว้ที่นี่ เพื่อ Reset ค่าทุกครั้งที่ขึ้น Sub-criteria ใหม่ (กันเลขบวม)
        sub_conf_pool = {} 

        for lv in range(1, 6):
            k = str(lv)
            info = lv_details.get(k, {}) or {}
            is_passed = bool(info.get("is_passed", False))
            level_key = f"{sub_id}_L{lv}"
            
            # 🎯 1. Resolve Evidence Sources
            sources = info.get("evidence_sources") or info.get("evidences") or []
            if not sources and level_key in global_evidence_map:
                ext_ev = global_evidence_map[level_key]
                sources = [ext_ev] if isinstance(ext_ev, dict) else ext_ev

            # 🎯 2. PDCA Sync
            req_phases = info.get("required_pdca_phases", []) or ["P"]
            actual_found_tags = set()

            # ตรวจสอบว่า sources เป็น list ก่อน loop
            current_sources = sources if isinstance(sources, list) else []

            for src in current_sources:
                if not isinstance(src, dict): continue

                fname = str(src.get("filename") or src.get("file") or src.get("source") or "เอกสารอ้างอิง")
                
                # 🚩 FIX: Confidence Extraction
                raw_val = None
                if "|SCORE:" in fname:
                    try: raw_val = float(fname.split("SCORE:")[-1])
                    except: pass
                
                if raw_val is None:
                    raw_val = src.get("rerank_score") or src.get("relevance_score") or src.get("confidence")

                try:
                    conf_val = float(raw_val) if raw_val is not None else 0.5
                    if 0 < conf_val <= 1.0: conf_val *= 100
                except: conf_val = 50.0
                
                clean_fname = fname.split("|")[0]
                # เก็บค่าความเชื่อมั่นสูงสุดรายไฟล์
                sub_conf_pool[clean_fname] = max(conf_val / 100, sub_conf_pool.get(clean_fname, 0))

                # 🚩 FIX: PDCA Tag ต้องดึงมาใส่ใน Object ที่จะส่งให้ UI
                tag = str(src.get("pdca_tag") or src.get("pdca") or "D").upper()
                if tag not in ["P", "D", "C", "A"]: tag = "D"
                actual_found_tags.add(tag)

                grouped_sources[k].append({
                    "filename": clean_fname,
                    "document_uuid": src.get("stable_doc_uuid") or src.get("doc_id"),
                    "page": str(src.get("page", "1")),
                    "pdca_tag": tag, # 🚩 จุดที่ UI นำไปแสดงผล
                    "confidence": round(conf_val, 1),
                    "text": src.get("content") or src.get("snippet") or "ไม่พบรายละเอียดข้อความ"
                })

            # คำนวณ Coverage
            actual_passed_phases = [p for p in req_phases if p in actual_found_tags]
            calc_percentage = (len(actual_passed_phases) / len(req_phases)) * 100 if req_phases else 0

            pdca_coverage[k] = {
                "percentage": round(calc_percentage, 1),
                "statement": info.get("rubric_statement") or "",
                "required_phases": req_phases,
                "actual_phases": list(actual_found_tags),
                "status": "PASS" if calc_percentage >= 100 else "GAP"
            }

            ui_levels[k] = {
                "level": lv, 
                "is_passed": is_passed,
                "score": round(float(info.get("score") or (1.0 if is_passed else 0.0)), 2),
                "reason": info.get("reason", "ไม่พบข้อมูลการประเมิน"),
                "coaching_insight": info.get("coaching_insight", ""),
                "action_plan": info.get("action_plan") or info.get("atomic_action_plan", [])
            }
    
            pdca_matrix.append({
                "level": lv, 
                "is_passed": is_passed,
                "pdca": {p: (1 if p in actual_found_tags else 0) for p in ["P", "D", "C", "A"]}
            })

        # --- [4] FINAL ASSEMBLY ---
        # 🚩 FIX: คำนวณ Traceability จาก Pool ของหัวข้อนี้ (เฉลี่ยไฟล์ที่ไม่ซ้ำ)
        if sub_conf_pool:
            avg_conf_total = sum(sub_conf_pool.values()) / len(sub_conf_pool)
            final_traceability = min(avg_conf_total * 100, 100)
        else:
            final_traceability = 0

        processed.append({
            "code": sub_id,
            "name": sub.get("sub_criteria_name", "Unknown"),
            "level": f"L{sub.get('highest_full_level', 0)}",
            "score": round(float(sub.get("weighted_score", 0.0)), 2),
            "strategic_focus": sub.get("strategic_focus", ""),
            "sub_roadmap": ui_sub_roadmap,
            "pdca_matrix": pdca_matrix,
            "pdca_coverage": pdca_coverage,
            "level_details": ui_levels,
            "grouped_sources": grouped_sources,
            "audit_confidence": {
                "source_count": len(sub_conf_pool),
                "traceability_score": round(final_traceability, 1)
            }
        })
        radar_data.append({"axis": sub_id, "value": sub.get("highest_full_level", 0)})

    try:
        processed.sort(key=lambda x: [int(p) for p in x["code"].split(".") if p.isdigit()])
    except: pass
    
    return {
        "status": summary.get("status", "COMPLETED"),
        "record_id": metadata.get("record_id") or raw_data.get("record_id"),
        "tenant": metadata.get("tenant", "n/a"),
        "year": metadata.get("year", "2567"),
        "enabler": metadata.get("enabler") or raw_data.get("enabler"),
        "level": str(summary.get("overall_max_level") or summary.get("maturity_level") or "0").replace("L", ""),
        "score": round(float(summary.get("total_weighted_score") or 0.0), 2),
        "enabler_roadmap": ui_global_roadmap,
        "radar_data": radar_data,
        "sub_criteria": processed
    }

# def _transform_result_for_ui(raw_data: Dict[str, Any], current_user: Any = None) -> Dict[str, Any]:
#     if not raw_data or not isinstance(raw_data, dict):
#         return {"status": "FAILED", "message": "Invalid data format"}

#     # --- [0] RESOLVE CORE DATA ---
#     res = raw_data.get("result") or raw_data.get("assessment_result") or raw_data
#     metadata = res.get("metadata", {})
#     summary = res.get("summary") or res.get("result_summary", {})
#     global_evidence_map = raw_data.get("evidence_map") or res.get("evidence_map") or {}

#     # --- [1] ENABLER ROADMAP ---
#     raw_global_roadmap = res.get("enabler_roadmap") or res.get("strategic_roadmap") or {}
#     ui_global_roadmap = {
#         "status": raw_global_roadmap.get("status", "SUCCESS"),
#         "overall_strategy": raw_global_roadmap.get("overall_strategy") or "มุ่งเน้นการยกระดับตามมาตรฐาน SE-AM",
#         "phases": raw_global_roadmap.get("phases") or []
#     }

#     # --- [2] SUB-CRITERIA PROCESSING ---
#     processed = []
#     radar_data = []
#     sub_list = res.get("sub_criteria_details") or res.get("sub_criteria_results") or []

#     for sub in sub_list:
#         sub_id = str(sub.get("sub_id", "N/A"))
#         lv_details = sub.get("level_details", {}) or {}
        
#         ui_sub_roadmap = {
#             "strategy": (sub.get("sub_roadmap") or {}).get("overall_strategy") or sub.get("strategic_focus", ""),
#             "phases": (sub.get("sub_roadmap") or {}).get("phases") or [],
#             "is_gap_detected": (sub.get("sub_roadmap") or {}).get("is_gap_detected", False)
#         }

#         ui_levels = {}
#         pdca_matrix = []
#         pdca_coverage = {}
#         grouped_sources = {str(i): [] for i in range(1, 6)}
        
#         # 🚩 สำหรับเก็บค่าสูงสุดของไฟล์เพื่อคำนวณ Traceability ท้ายสุด
#         sub_conf_pool = {} 

#         for lv in range(1, 6):
#             k = str(lv)
#             info = lv_details.get(k, {}) or {}
#             is_passed = bool(info.get("is_passed", False))
#             level_key = f"{sub_id}_L{lv}"
            
#             # 🎯 1. Resolve Evidence Sources
#             sources = info.get("evidence_sources") or info.get("evidences") or []
#             if not sources and level_key in global_evidence_map:
#                 ext_ev = global_evidence_map[level_key]
#                 sources = [ext_ev] if isinstance(ext_ev, dict) else ext_ev

#             # 🎯 2. PDCA Sync - สร้าง Set เก็บข้อมูลจริงที่เจอใน Level นี้
#             req_phases = info.get("required_pdca_phases", []) or ["P"]
#             actual_found_tags = set()

#             for src in sources:
#                 # Resolve Filename & Score
#                 fname = str(src.get("filename") or src.get("file") or src.get("source") or "เอกสารอ้างอิง")
                
#                 # 🚩 FIX ISSUE 3: Confidence Extraction (ดึงจาก SCORE:0.xxxx)
#                 raw_val = None
#                 if "|SCORE:" in fname:
#                     try:
#                         raw_val = float(fname.split("SCORE:")[-1])
#                     except: pass
                
#                 # ถ้าในชื่อไฟล์ไม่มี ให้ไปดูที่ Key มาตรฐาน
#                 if raw_val is None:
#                     raw_val = src.get("rerank_score") or src.get("relevance_score") or src.get("confidence")

#                 try:
#                     conf_val = float(raw_val) if raw_val is not None else 0.5
#                     # ถ้ามาเป็น 0.75 ให้คูณ 100 เป็น 75.0
#                     if 0 < conf_val <= 1.0: conf_val *= 100
#                 except:
#                     conf_val = 50.0
                
#                 clean_fname = fname.split("|")[0]
#                 # เก็บค่าความเชื่อมั่นสูงสุดรายไฟล์ (Scale 0-1)
#                 sub_conf_pool[clean_fname] = max(conf_val / 100, sub_conf_pool.get(clean_fname, 0))

#                 # สกัด PDCA Tag
#                 tag = str(src.get("pdca_tag") or src.get("pdca") or "D").upper()
#                 if tag not in ["P", "D", "C", "A"]: tag = "D"
#                 actual_found_tags.add(tag)

#                 grouped_sources[k].append({
#                     "filename": clean_fname,
#                     "document_uuid": src.get("stable_doc_uuid") or src.get("doc_id"),
#                     "page": str(src.get("page", "1")),
#                     "pdca_tag": tag,
#                     "confidence": round(conf_val, 1),
#                     "text": src.get("content") or src.get("snippet") or "ไม่พบรายละเอียดข้อความ"
#                 })

#             # คำนวณ Coverage
#             actual_passed_phases = [p for p in req_phases if p in actual_found_tags]
#             calc_percentage = (len(actual_passed_phases) / len(req_phases)) * 100 if req_phases else 0

#             pdca_coverage[k] = {
#                 "percentage": round(calc_percentage, 1),
#                 "statement": info.get("rubric_statement") or "",
#                 "required_phases": req_phases,
#                 "actual_phases": list(actual_found_tags),
#                 "status": "PASS" if calc_percentage >= 100 else "GAP"
#             }

#             ui_levels[k] = {
#                 "level": lv, 
#                 "is_passed": is_passed,
#                 "score": round(float(info.get("score") or (1.0 if is_passed else 0.0)), 2),
#                 "reason": info.get("reason", "ไม่พบข้อมูลการประเมิน"),
#                 "coaching_insight": info.get("coaching_insight", ""),
#                 "action_plan": info.get("action_plan") or info.get("atomic_action_plan", [])
#             }
    
#             # 🚩 FIX ISSUE 1: PDCA Matrix แสดงผลที่แท้จริง (ไม่ใช่ Required)
#             pdca_matrix.append({
#                 "level": lv, 
#                 "is_passed": is_passed,
#                 "pdca": {p: (1 if p in actual_found_tags else 0) for p in ["P", "D", "C", "A"]}
#             })

#         # --- [4] FINAL ASSEMBLY ---
#         # 🚩 FIX ISSUE 2: Traceability Score (เฉลี่ยความเชื่อมั่นไฟล์ที่ไม่ซ้ำกัน)

#         if sub_conf_pool:
#             avg_conf_total = sum(sub_conf_pool.values()) / len(sub_conf_pool)
#             final_traceability = min(avg_conf_total * 100, 100)
#         else:
#             final_traceability = 0

#         processed.append({
#             "code": sub_id,
#             "name": sub.get("sub_criteria_name", "Unknown"),
#             "level": f"L{sub.get('highest_full_level', 0)}",
#             "score": round(float(sub.get("weighted_score", 0.0)), 2),
#             "strategic_focus": sub.get("strategic_focus", ""),
#             "sub_roadmap": ui_sub_roadmap,
#             "pdca_matrix": pdca_matrix,
#             "pdca_coverage": pdca_coverage,
#             "level_details": ui_levels,
#             "grouped_sources": grouped_sources,
#             "audit_confidence": {
#                 "source_count": len(sub_conf_pool),
#                 "traceability_score": round(final_traceability, 1)
#             }
#         })
#         radar_data.append({"axis": sub_id, "value": sub.get("highest_full_level", 0)})

#     try:
#         processed.sort(key=lambda x: [int(p) for p in x["code"].split(".") if p.isdigit()])
#     except: pass
    
#     return {
#         "status": summary.get("status", "COMPLETED"),
#         "record_id": metadata.get("record_id") or raw_data.get("record_id"),
#         "tenant": metadata.get("tenant", "n/a"),
#         "year": metadata.get("year", "2567"),
#         "enabler": metadata.get("enabler") or raw_data.get("enabler"),
#         "level": str(summary.get("overall_max_level") or summary.get("maturity_level") or "0").replace("L", ""),
#         "score": round(float(summary.get("total_weighted_score") or 0.0), 2),
#         "enabler_roadmap": ui_global_roadmap,
#         "radar_data": radar_data,
#         "sub_criteria": processed
#     }

def set_thai_font(run, size=14, bold=False, color=None):
    """ตั้งค่าฟอนต์ TH Sarabun New ให้รองรับทั้งภาษาไทยและอังกฤษ"""
    run.font.name = 'TH Sarabun New'
    # บังคับให้ XML ภายในใช้ TH Sarabun New สำหรับอักขระพิเศษและภาษาไทย
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'TH Sarabun New')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'TH Sarabun New')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def set_cell_background(cell, fill_color):
    """ระบายสีพื้นหลังให้ Cell ในตาราง (fill_color คือ hex code เช่น 'D9EAD3')"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_docx_report_similar_to_ui(ui_data: dict) -> Document:
    """
    [v2026.FINAL - Revised for level_details structure]
    สร้างรายงาน Word โดยอิงข้อมูลจาก UI-Ready JSON ที่ผ่านการ transform มาแล้ว
    """
    doc = Document()
    
    # 1. หัวข้อรายงาน (Header)
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = header.add_run(f"รายงานผลการประเมิน Maturity Audit ({ui_data.get('enabler', 'KM')})\n")
    set_thai_font(run_h, size=20, bold=True, color=RGBColor(30, 58, 138))

    # 2. ส่วนสรุปภาพรวม (Overall Summary)
    maturity_lv = str(ui_data.get('level', '0'))
    total_score = ui_data.get('score', 0)
    full_score = ui_data.get('full_score', 5)
    metrics = ui_data.get('metrics', {})

    sum_p = doc.add_paragraph()
    sum_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sum = sum_p.add_run(
        f"ระดับวุฒิภาวะ: L{maturity_lv} | คะแนนรวม: {total_score}/{full_score} "
        f"({metrics.get('completion_rate', 0)}% ผ่านเกณฑ์ย่อย)"
    )
    set_thai_font(run_sum, size=16, bold=True, color=RGBColor(22, 101, 52))

    # 3. วนลูปรายหัวข้อเกณฑ์ย่อย (Sub-Criteria)
    for item in ui_data.get('sub_criteria', []):
        # หัวข้อเกณฑ์
        doc.add_paragraph() 
        title_p = doc.add_paragraph()
        run_title = title_p.add_run(f"เกณฑ์ {item.get('code', '')}: {item.get('name', '')}")
        set_thai_font(run_title, size=16, bold=True, color=RGBColor(30, 58, 138))

        # --- 3.1 ตาราง Audit Confidence ---
        conf_table = doc.add_table(rows=1, cols=3)
        conf_table.style = 'Table Grid'
        conf = item.get('audit_confidence', {})
        
        # จัดการค่าความมั่นใจ (Traceability)
        trace_val = conf.get('traceability_score', 0)
        if trace_val <= 1.0: trace_val = int(trace_val * 100) # แปลง 0.8 -> 80
        
        metrics_cells = [
            ("Independence", f"{conf.get('source_count', 0)} Files"),
            ("Traceability", f"{trace_val}% Confidence"),
            ("Audit Status", f"{conf.get('level', 'VERIFIED')}")
        ]
        
        for i, (label, val) in enumerate(metrics_cells):
            cell = conf_table.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_thai_font(p.add_run(label), size=10, bold=True)
            set_thai_font(p.add_run(f"\n{val}"), size=12, bold=True)
            set_cell_background(cell, "F3F4F6")

        # --- 3.2 PDCA Capability Matrix ---
        doc.add_paragraph()
        set_thai_font(doc.add_paragraph().add_run("📊 PDCA Capability Matrix:"), size=13, bold=True)
        pdca_table = doc.add_table(rows=2, cols=5)
        pdca_table.style = 'Table Grid'
        
        # ดึงข้อมูลจาก pdca_matrix ที่เรา transform มาแล้ว
        for i, lv_data in enumerate(item.get('pdca_matrix', [])):
            if i >= 5: break
            cell_top = pdca_table.cell(0, i)
            p_top = cell_top.paragraphs[0]
            p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_thai_font(p_top.add_run(f"Level {lv_data['level']}"), bold=True)
            
            if lv_data.get('is_passed'):
                set_cell_background(cell_top, "D9EAD3") # สีเขียวถ้าผ่าน

            cell_bot = pdca_table.cell(1, i)
            p_bot = cell_bot.paragraphs[0]
            p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # แสดง P D C A แยกตามสถานะ (1=เขียว, 0=แดง)
            for char, val in lv_data.get('pdca', {}).items():
                run_char = p_bot.add_run(f" {char} ")
                color = RGBColor(22, 101, 52) if val == 1 else RGBColor(185, 28, 28)
                set_thai_font(run_char, size=11, bold=True, color=color)

        # --- 3.3 รายการหลักฐาน (Evidence Mapping) ---
        doc.add_paragraph()
        set_thai_font(doc.add_paragraph().add_run("📎 รายการหลักฐานที่ตรวจพบ (Evidence Mapping):"), size=12, bold=True)
        
        grouped_sources = item.get('grouped_sources', {})
        has_evidence = False
        
        # วนลูป Level 1-5 เพื่อเรียงลำดับหลักฐาน
        for lv_key in ["1", "2", "3", "4", "5"]:
            sources = grouped_sources.get(lv_key, [])
            for src in sources:
                has_evidence = True
                # ทำความสะอาดชื่อไฟล์ (เอา SCORE ออกถ้ามี)
                clean_filename = src.get('filename', '').split('|')[0]
                evi_text = (
                    f"Level {lv_key}: {clean_filename} "
                    f"(หน้า {src.get('page', '1')}) - "
                    f"Relevance: {src.get('rerank_score', 0)}%"
                )
                p_evi = doc.add_paragraph(style='List Bullet')
                set_thai_font(p_evi.add_run(evi_text), size=10)
        
        if not has_evidence:
            set_thai_font(doc.add_paragraph().add_run("- ไม่พบหลักฐานแนบในหัวข้อนี้ -"), size=10)

        # --- 3.4 Insights & Recommendations ---
        # Strength
        doc.add_paragraph()
        set_thai_font(doc.add_paragraph().add_run("💡 AI Strength Summary:"), size=13, bold=True, color=RGBColor(22, 101, 52))
        reason_txt = item.get('reason', 'ผ่านเกณฑ์ตามมาตรฐานที่กำหนด')
        set_thai_font(doc.add_paragraph(reason_txt).runs[0], size=12)

        # Next Step
        set_thai_font(doc.add_paragraph().add_run("🚀 Next Step Recommendation:"), size=13, bold=True, color=RGBColor(30, 58, 138))
        next_step_txt = item.get('next_step', 'รักษามาตรฐานการดำเนินงานและเตรียมความพร้อมสู่ระดับถัดไป')
        set_thai_font(doc.add_paragraph(next_step_txt).runs[0], size=12)

        doc.add_page_break() 
        
    return doc

@assessment_router.get("/status/{record_id}")
async def get_assessment_status(
    record_id: str, 
    current_user: UserMe = Depends(get_current_user)
):
    """
    [v2026.01.27 — THE SHIELDED STATUS REVISE]
    - ⚡ Layer 1: Check Database (The Truth) -> ป้องกัน Race Condition หลัง Start
    - 🧠 Layer 2: Check Active Tasks (Memory) -> สำหรับ Real-time Update
    - 📂 Layer 3: Check Disk (Persistence) -> สำหรับงานที่จบไปแล้ว
    """
    db = SessionLocal()
    try:
        # --- LAYER 1: CHECK DATABASE (แก้ปัญหา 404 หลังกด Start) ---
        task_record = db.query(AssessmentTaskTable).filter(
            AssessmentTaskTable.record_id == record_id
        ).first()

        # ถ้าเจอใน DB แปลว่างานถูกลงทะเบียนแล้วแน่นอน (ไม่ควรตอบ 404)
        if task_record:
            # ถ้าสถานะยังไม่เสร็จ ให้ตอบสถานะจาก DB/Memory
            if task_record.status not in ["COMPLETED", "FAILED"]:
                # เช็ค Memory ประกอบ (ถ้ามีข้อมูลที่ละเอียดกว่า)
                active_tasks = globals().get("ACTIVE_TASKS", {})
                mem_task = active_tasks.get(record_id, {})
                
                return {
                    "status": task_record.status, # QUEUED, PROCESSING
                    "record_id": record_id,
                    "progress": mem_task.get("progress") or task_record.progress_percent or 5,
                    "message": mem_task.get("message") or task_record.progress_message or "กำลังเตรียมการ...",
                    "enabler": task_record.enabler,
                    "is_final": False,
                    "updated_at": datetime.now().isoformat()
                }
            
            # ถ้าใน DB บอกว่า FAILED ให้แจ้งทันที
            if task_record.status == "FAILED":
                return {
                    "status": "FAILED",
                    "record_id": record_id,
                    "message": task_record.progress_message or "การประเมินล้มเหลว",
                    "is_final": True
                }

        # --- LAYER 2: CHECK DISK (สำหรับงานที่ COMPLETED แล้ว) ---
        # ฟังก์ชัน _find_assessment_file จะทำ Deep Scan หาไฟล์ JSON
        try:
            file_path = _find_assessment_file(record_id, current_user)
        except HTTPException:
            file_path = None

        if not file_path or not os.path.exists(file_path):
            # กรณีที่ไม่เจอทั้งใน DB และ Disk จริงๆ
            if not task_record:
                logger.warning(f"🔍 [Status] Record not found anywhere: {record_id}")
                return {
                    "status": "NOT_FOUND",
                    "record_id": record_id,
                    "message": "ไม่พบรหัสการประเมินนี้ในระบบ",
                    "suggestion": "กรุณาตรวจสอบรหัสอีกครั้ง หรือเริ่มการประเมินใหม่"
                }
            
            # กรณีมีใน DB ว่าเสร็จแล้วแต่ไฟล์ยังไม่มา (I/O Delay)
            return {
                "status": "PROCESSING",
                "record_id": record_id,
                "progress": 95,
                "message": "AI ประมวลผลเสร็จแล้ว กำลังบันทึกข้อมูลลงระบบไฟล์...",
                "is_final": False
            }

        # --- LAYER 3: DATA TRANSFORMATION ---
        with open(file_path, "r", encoding="utf-8") as f:
            raw_data = json.load(f)

        # ตรวจสอบสิทธิ์จาก Data ในไฟล์
        meta = raw_data.get("metadata", {}) or raw_data.get("summary", {}) or {}
        check_user_permission(current_user, meta.get("tenant"), meta.get("enabler"))

        # Transform ข้อมูลให้ UI
        ui_result = _transform_result_for_ui(raw_data, current_user)
        ui_result["status"] = "COMPLETED"
        ui_result["record_id"] = record_id
        ui_result["is_final"] = True

        return ui_result

    except Exception as e:
        logger.error(f"💥 [Status Error] {record_id}: {str(e)}", exc_info=True)
        if isinstance(e, HTTPException): raise e
        raise HTTPException(status_code=500, detail="Internal Server Error ในการดึงสถานะ")
    finally:
        db.close()

@assessment_router.get("/history")
async def get_assessment_history(
    tenant: str, 
    year: Optional[str] = Query(None),
    enabler: Optional[str] = Query(None),
    current_user: UserMe = Depends(get_current_user)
):
    """
    [v2026.FINAL.HISTORY] - ดึงข้อมูลประวัติการประเมิน
    - แก้ไขปัญหา Scope เป็น ALL โดยดึงจาก metadata.sub_id หรือรายการเกณฑ์จริง
    - แก้ไขปัญหา Level ไม่แสดง โดยดึงจาก result_summary.maturity_level
    - รองรับโครงสร้างไฟล์ Hybrid ทั้งเก่าและใหม่
    """
    check_user_permission(current_user, tenant)
    history_list = []
    from config.global_vars import DATA_STORE_ROOT
    from datetime import datetime
    
    norm_tenant = _n(tenant)
    search_roots = [
        os.path.join(DATA_STORE_ROOT, norm_tenant, "exports"),
        os.path.join("data_store", norm_tenant, "exports")
    ]
    
    tenant_export_root = next((p for p in search_roots if os.path.exists(p)), None)
    if not tenant_export_root:
        return {"items": [], "total_found": 0, "message": "No export data found"}

    user_allowed_enablers = [e.upper() for e in current_user.enablers]
    target_enabler = enabler.upper() if enabler else None

    # จัดการเรื่องปีงบประมาณ
    if not year or str(year).lower() == "all":
        search_years = [d for d in os.listdir(tenant_export_root) if d.isdigit()]
    else:
        search_years = [str(year)]

    for y in search_years:
        year_path = os.path.join(tenant_export_root, y)
        if not os.path.exists(year_path): continue

        for root, _, files in os.walk(year_path):
            for f in files:
                if not f.lower().endswith(".json"): continue
                file_path = os.path.join(root, f)
                
                try:
                    with open(file_path, "r", encoding="utf-8") as jf:
                        data = json.load(jf)

                    # 1. แยกส่วนข้อมูล (v2026 ใช้ metadata และ result_summary)
                    metadata = data.get("metadata", {})
                    res_sum = data.get("result_summary", {})
                    old_sum = data.get("summary", {})

                    # 2. RECORD ID
                    record_id = data.get("record_id") or metadata.get("record_id") or old_sum.get("record_id")
                    if not record_id:
                        parts = f.replace(".json", "").split("_")
                        record_id = parts[2] if len(parts) >= 3 else f.replace(".json", "")

                    # 3. ENABLER
                    file_enabler = (metadata.get("enabler") or res_sum.get("enabler") or "KM").upper()
                    if file_enabler not in user_allowed_enablers: continue
                    if target_enabler and file_enabler != target_enabler: continue

                    # 4. SCOPE (แก้ไขเพื่อให้ดึง sub_id จากโครงสร้างใหม่ได้แม่นยำ)
                    details = data.get("sub_criteria_details", [])
                    found_subs = []

                    # ดึงรายการ sub_id ที่มีการประเมินจริงในไฟล์นี้
                    if isinstance(details, list):
                        for detail in details:
                            # กรณีรันผ่าน Worker (โครงสร้าง Tier-2)
                            sub_results = detail.get("sub_criteria_results", [])
                            for res in sub_results:
                                sid = res.get("sub_id")
                                if sid: found_subs.append(str(sid))
                            
                            # กรณีไฟล์โครงสร้างเดี่ยว (Tier-1)
                            if not sub_results and detail.get("sub_id"):
                                found_subs.append(str(detail.get("sub_id")))

                    # กำจัดตัวซ้ำและเรียงลำดับ (เช่น 1.1, 1.2)
                    unique_subs = sorted(list(set(found_subs)))

                    # --- ตัดสินใจเลือกข้อความที่จะแสดงบน UI ---
                    if len(unique_subs) == 1:
                        # 🎯 กรณีประเมินรายข้อ: แสดงเลขข้อไปเลย เช่น "1.1"
                        scope = unique_subs[0]
                    elif 1 < len(unique_subs) <= 3:
                        # 🎯 กรณีประเมินบางส่วน (2-3 ข้อ): แสดงเลขข้อเรียงกัน เช่น "1.1, 1.2"
                        scope = ", ".join(unique_subs)
                    elif len(unique_subs) > 3:
                        # 🎯 กรณีประเมินจำนวนมาก แต่ไม่ครบทั้ง Enabler
                        scope = "MULTI"
                    else:
                        # 🎯 Fallback สุดท้ายถ้าหาไม่เจอจริงๆ หรือเป็นการประเมินทั้ง Enabler
                        raw_scope = metadata.get("sub_id") or old_sum.get("sub_criteria_id") or "ALL"
                        scope = str(raw_scope).upper()

                    scope = scope.upper()

                    # 5. LEVEL LOGIC
                    display_level = res_sum.get("maturity_level") or old_sum.get("highest_pass_level")
                    if display_level:
                        l_str = str(display_level).strip().upper()
                        display_level = l_str if l_str.startswith("L") else f"L{l_str}"
                    else:
                        display_level = "N/A"

                    # 6. SCORE
                    total_score = round(safe_float(
                        res_sum.get("total_weighted_score") or 
                        old_sum.get("total_weighted_score") or 0.0
                    ), 2)

                    # 7. DATE PARSING
                    date_candidates = [
                        metadata.get("exported_at"), # มาตรฐานใหม่
                        metadata.get("export_at"),
                        old_sum.get("timestamp")
                    ]
                    date_str, parsed_dt = "N/A", None
                    for cand in date_candidates:
                        if cand:
                            try:
                                parsed_dt = datetime.fromisoformat(str(cand).replace('Z', '+00:00'))
                                date_str = parsed_dt.isoformat()
                                break
                            except: continue

                    if not parsed_dt:
                        mtime = os.path.getmtime(file_path)
                        parsed_dt = datetime.fromtimestamp(mtime)
                        date_str = parsed_dt.isoformat()

                    history_list.append({
                        "record_id": record_id,
                        "date": date_str,
                        "date_dt": parsed_dt,
                        "tenant": tenant,
                        "year": y,
                        "enabler": file_enabler,
                        "scope": scope,
                        "level": display_level,
                        "score": total_score,
                        "status": "COMPLETED"
                    })

                except Exception as e:
                    logger.error(f"❌ Skip corrupted/old file {f}: {e}")
                    continue

    # 8. Sort ผลลัพธ์ตามเวลาล่าสุด
    sorted_history = sorted(history_list, key=lambda x: x['date_dt'] or datetime.min, reverse=True)
    for item in sorted_history: item.pop('date_dt', None)

    return {
        "items": sorted_history,
        "total_found": len(history_list),
        "displayed": len(sorted_history)
    }

# ------------------------------------------------------------------
# 1. Start Assessment Endpoint
# ------------------------------------------------------------------
@assessment_router.post("/start")
async def start_assessment(
    request: StartAssessmentRequest, 
    background_tasks: BackgroundTasks, 
    current_user: UserMe = Depends(get_current_user)
):
    """
    [FINAL v2026.6.20 — Start Assessment + Friendly Response]
    - Permission + Data Integrity Check
    - Persistent Task Entry (DB)
    - Background Worker Delegation
    - Response เป็นมิตร + estimated_time
    """
    enabler_uc = request.enabler.upper()
    target_year = str(request.year if request.year else (current_user.year or DEFAULT_YEAR)).strip()
    target_sub = str(request.sub_criteria).strip().lower() if request.sub_criteria else "all"

    # 1. Permission & Data Integrity Check
    check_user_permission(current_user, request.tenant, enabler_uc)

    vs_path = get_vectorstore_collection_path(request.tenant, target_year, "evidence", enabler_uc)
    if not os.path.exists(vs_path):
        raise HTTPException(status_code=400, detail=f"Data Store สำหรับ {enabler_uc}/{target_year} ยังไม่ถูกสร้าง")

    # 2. Generate Traceable Record ID
    record_id = uuid.uuid4().hex[:12]
    
    # 3. Persistent Task Entry
    db = SessionLocal()
    try:
        new_task = AssessmentTaskTable(
            record_id=record_id,
            user_id=current_user.id,
            tenant=request.tenant,
            year=target_year,
            enabler=enabler_uc,
            sub_criteria=target_sub,
            status="QUEUED",
            progress_percent=5,
            progress_message="กำลังคิวงานประเมิน..."
        )
        db.add(new_task)
        db.commit()
        db.refresh(new_task)
    except Exception as e:
        logger.error(f"❌ Initial DB Error: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="ไม่สามารถลงทะเบียนงานประเมินได้ กรุณาลองใหม่")
    finally:
        db.close()

    # 4. Delegate to Background Worker
    background_tasks.add_task(
        run_assessment_engine_task,
        assessment_semaphore,     # ส่งตัวแปรไปเลย ไม่ต้องระบุชื่อ
        record_id,
        request.tenant,
        target_year,
        enabler_uc,
        target_sub,
        request.sequential_mode
    )

    return {
        "record_id": record_id,
        "status": "QUEUED",
        "message": f"เริ่มการประเมิน {enabler_uc} เรียบร้อยแล้ว (กำลังคิวงาน)",
        "estimated_time": "20-40 นาที (ขึ้นกับจำนวนเอกสารและโหมด sequential)",
        "poll_url": f"/api/assess/status/{record_id}",
        "poll_interval_seconds": 15
    }

# ------------------------------------------------------------------
# 2. Background Task Engine (Robust Implementation)
# ------------------------------------------------------------------
async def run_assessment_engine_task(
    semaphore: asyncio.Semaphore, 
    record_id: str, 
    tenant: str, 
    year: str, 
    enabler: str, 
    sub_id: str, 
    sequential: bool
):
    """
    [v2026.FINAL.REVISED — Robust Background Worker]
    - 🔒 Semaphore Control: จำกัดการใช้ GPU/RAM พร้อมกัน
    - 🧹 Memory Management: Cleanup ทรัพยากรทันทีหลังจบงาน
    - 📊 Accurate DB Sync: บันทึก Scope และ Level จริงลงฐานข้อมูล
    """
    # 1. 🔒 Acquire Semaphore (Queue Management)
    async with semaphore:
        logger.info(f"⚙️ [Task {record_id}] Processing Started (Semaphore Acquired)...")
        engine = None
        vsm = None
        
        try:
            # --- [STEP 1: RESOURCE HYDRATION] ---
            db_update_task_status(record_id, 10, "เชื่อมต่อ Vector Database และโหลด Mapping...")
            
            # รัน CPU/IO Bound tasks ใน Thread
            vsm = await asyncio.to_thread(
                load_all_vectorstores, tenant, year, None, EVIDENCE_DOC_TYPES, enabler
            )
            doc_map_raw = await asyncio.to_thread(
                load_doc_id_mapping, EVIDENCE_DOC_TYPES, tenant, year, enabler
            )
            doc_map = {d_id: d.get("file_name", d_id) for d_id, d in doc_map_raw.items()}

            # --- [STEP 2: ENGINE & MODEL SETUP] ---
            db_update_task_status(record_id, 20, f"โหลด AI Model ({DEFAULT_LLM_MODEL_NAME})...")
            
            llm = await asyncio.to_thread(
                create_llm_instance, model_name=DEFAULT_LLM_MODEL_NAME, temperature=0.0
            )
            
            config = AssessmentConfig(
                enabler=enabler, 
                tenant=tenant, 
                year=year, 
                force_sequential=sequential
            )
            
            engine = SEAMPDCAEngine(
                config=config, 
                llm_instance=llm, 
                logger_instance=logger, 
                doc_type=EVIDENCE_DOC_TYPES, 
                vectorstore_manager=vsm, 
                document_map=doc_map,
                record_id=record_id 
            )

            # --- [STEP 3: CORE ASSESSMENT] ---
            db_update_task_status(record_id, 35, "AI กำลังตรวจสอบหลักฐาน (RAG Assessment)...")
            
            result = await asyncio.to_thread(
                engine.run_assessment, 
                target_sub_id=sub_id, 
                export=True, 
                record_id=record_id,
                vectorstore_manager=vsm,
                sequential=sequential
            )

            # --- [STEP 4: FINALIZE & SYNC DB] ---
            if isinstance(result, dict) and result.get("status") == "FAILED":
                error_msg = result.get("error_message", "AI Engine Error")
                db_update_task_status(record_id, 0, f"ล้มเหลว: {error_msg}", status="FAILED")
            else:
                # 🎯 [CRITICAL] บันทึกผลลัพธ์สุดท้ายและสรุปคะแนนลง DB ทันที
                # เพื่อให้ API /history ดึง Scope และ Level ที่ถูกต้องมาแสดงได้
                await asyncio.to_thread(db_finish_task, record_id, result)
                
                # อัปเดตสถานะสุดท้าย
                db_update_task_status(record_id, 100, "การประเมินเสร็จสมบูรณ์", status="COMPLETED")
                logger.info(f"✅ [Task {record_id}] Finished Successfully")
                
        except Exception as e:
            logger.error(f"💥 [Task {record_id}] Critical Failure: {str(e)}", exc_info=True)
            db_update_task_status(record_id, 0, f"ระบบขัดข้อง: {str(e)}", status="FAILED")
            
        finally:
            # --- [STEP 5: POLISHING & CLEANUP] ---
            # สำคัญมากสำหรับการรันบน GPU L40S เพื่อไม่ให้ Memory ค้าง
            import gc
            import torch
            
            # ลบ Instance ขนาดใหญ่
            del engine
            del vsm
            
            if torch.cuda.is_available():
                torch.cuda.empty_cache() # คืน VRAM
            
            gc.collect() # คืน RAM
            logger.info(f"🧹 [Task {record_id}] Memory cleanup completed.")
            

def _find_assessment_file(search_id: str, current_user: UserMe) -> str:
    """
    [HYBRID DEEP SEARCH v2026.3]
    - ชั้น 1: DB Hit (ค้นจาก Database ตรงๆ)
    - ชั้น 2: Fast Disk Scan (ค้นจากชื่อไฟล์ - ถ้ามี ID ในชื่อ)
    - ชั้น 3: Deep Disk Scan (เปิดอ่าน JSON Metadata เพื่อหา record_id) **NEW**
    """
    norm_tenant = _n(current_user.tenant)
    norm_search = str(search_id).strip().lower()

    # --- ชั้น 1: DB Hit (ค้นจาก SQLite) ---
    db = SessionLocal()
    try:
        res_record = db.query(AssessmentResultTable).filter(
            AssessmentResultTable.record_id == search_id
        ).first()
        
        if res_record and res_record.full_result_json:
            try:
                data = json.loads(res_record.full_result_json)
                # พยายามดึง path จากหลายที่ใน JSON
                db_path = data.get("export_path_used") or data.get("metadata", {}).get("full_path")
                if db_path and os.path.exists(db_path):
                    logger.info(f"⚡ [Search] DB Hit! Found: {db_path}")
                    return db_path
            except: pass
    finally:
        db.close()

    # --- ชั้น 2 & 3: Disk Scan (Fallback) ---
    search_paths = [
        os.path.join(DATA_STORE_ROOT, norm_tenant, "exports"),
        os.path.join("data_store", norm_tenant, "exports")
    ]
    
    logger.info(f"🔍 [Search] DB Miss. Deep Scanning Disk for ID: {norm_search}...")

    for s_path in search_paths:
        if not os.path.exists(s_path): continue
            
        for root, _, files in os.walk(s_path):
            for f in files:
                if not f.lower().endswith(".json"): continue
                
                full_path = os.path.join(root, f)
                
                # [Fast Scan] ถ้าโชคดีมี ID ในชื่อไฟล์
                if norm_search in f.lower():
                    logger.info(f"✅ [Search] Fast Scan Success: {full_path}")
                    return full_path
                
                # [Deep Scan] เปิดอ่าน Metadata ข้างใน (แก้ปัญหาชื่อไฟล์ไม่มี ID)
                try:
                    with open(full_path, "r", encoding="utf-8") as jf:
                        # อ่านแค่หัวไฟล์ (ป้องกันไฟล์ใหญ่แล้วค้าง)
                        first_part = jf.read(1000) 
                        # เช็คเบื้องต้นด้วย string search ก่อนโหลด json เต็ม
                        if norm_search in first_part:
                            jf.seek(0)
                            data = json.load(jf)
                            f_id = data.get("record_id") or data.get("metadata", {}).get("record_id")
                            if str(f_id).lower() == norm_search:
                                logger.info(f"🎯 [Search] Deep Scan Success! Found ID in Metadata: {full_path}")
                                return full_path
                except:
                    continue
                    
    logger.error(f"❌ [Search] Total Failure for ID: {norm_search}")
    raise HTTPException(
        status_code=404, 
        detail=f"ไม่พบไฟล์ผลการประเมิน (ID: {search_id}) แม้จะทำการ Deep Scan แล้ว"
    )

# ------------------------------------------------------------------
# 3. Task List API (สำหรับหน้า UI ดึงรายการมาโชว์)
# ------------------------------------------------------------------
@assessment_router.get("/tasks")
async def get_assessment_tasks(current_user: UserMe = Depends(get_current_user)):
    db = SessionLocal()
    try:
        # ดึงเฉพาะของ Tenant ตัวเอง เรียงลำดับใหม่สุดขึ้นก่อน
        tasks = db.query(AssessmentTaskTable).filter(
            AssessmentTaskTable.tenant == current_user.tenant
        ).order_by(AssessmentTaskTable.created_at.desc()).limit(20).all()
        
        return {"tasks": tasks}
    finally:
        db.close()


# ------------------------------------------------------------------
# 4. Download API (Full Revised)
# ------------------------------------------------------------------
@assessment_router.get("/download/{record_id}/{file_type}")
async def download_assessment_file(
    record_id: str,
    file_type: str,
    background_tasks: BackgroundTasks,
    current_user: UserMe = Depends(get_current_user)
):
    """
    [v2026.6 - Final Production]
    - แก้ไขการดึง Enabler จาก Metadata ใหม่
    - รองรับการสร้าง Word จาก UI Data ตัวล่าสุด
    """
    logger.info(f"📥 Download request: record_id={record_id}, type={file_type} by {current_user.email}")

    # 1. ค้นหาไฟล์ต้นฉบับ (JSON) ด้วย Deep Search ที่เราแก้กันก่อนหน้า
    json_path = _find_assessment_file(record_id, current_user)

    # 2. อ่านข้อมูล
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            raw_data = json.load(f)
    except Exception as e:
        logger.error(f"Error reading JSON: {e}")
        raise HTTPException(status_code=500, detail="ไม่สามารถอ่านไฟล์ข้อมูลได้")

    # 3. 🛡️ ตรวจสอบสิทธิ์ (ปรับตามโครงสร้างใหม่ v2026)
    # ดึงจาก metadata.enabler หรือ result_summary.enabler
    metadata = raw_data.get("metadata", {})
    res_sum = raw_data.get("result_summary", {})
    enabler = (metadata.get("enabler") or res_sum.get("enabler") or "KM").upper()
    
    check_user_permission(current_user, current_user.tenant, enabler)

    file_type = file_type.lower()

    # --- กรณีขอไฟล์ JSON ---
    if file_type == "json":
        return FileResponse(
            path=json_path,
            filename=f"SEAM_Result_{enabler}_{record_id}.json",
            media_type="application/json"
        )

    # --- กรณีขอไฟล์ Word (DOCX) ---
    elif file_type in ["word", "docx"]:
        logger.info(f"📄 Generating Word report for {record_id}...")

        # แปลงข้อมูลด้วย Transformer ตัวล่าสุด (v2026.5) เพื่อให้ได้ Roadmap และ Evidence ครบ
        ui_data = _transform_result_for_ui(raw_data)
        
        try:
            # ใช้ฟังก์ชันสร้าง Report (มั่นใจว่าส่ง ui_data ที่มี Roadmap ไปแล้ว)
            doc = create_docx_report_similar_to_ui(ui_data)
        except Exception as e:
            logger.error(f"Word Generation Error: {e}")
            raise HTTPException(status_code=501, detail=f"ระบบสร้างไฟล์ Word ขัดข้อง: {str(e)}")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            temp_path = tmp.name

        background_tasks.add_task(os.remove, temp_path)

        return FileResponse(
            path=temp_path,
            filename=f"SEAM_Report_{enabler}_{record_id}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        raise HTTPException(status_code=400, detail="รูปแบบไฟล์ไม่ถูกต้อง (รองรับ json, word)")

@assessment_router.get("/view-evidence/{record_id}/{lv}/{filename}")
async def view_evidence_file(
    record_id: str,
    lv: str,
    filename: str,
    current_user: UserMe = Depends(get_current_user)
):
    # 1. ค้นหาไฟล์ JSON ของ record นี้เพื่อดูว่ามีไฟล์นี้อ้างอิงอยู่จริงไหม (Security Check)
    json_path = _find_assessment_file(record_id, current_user)
    
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    # 2. ตรวจสอบสิทธิ์ Tenant/Enabler (ใช้ Logic เดิมที่คุณมี)
    metadata = data.get("metadata", {})
    check_user_permission(current_user, metadata.get("tenant"), metadata.get("enabler"))

    # 3. ประกอบ Path ไปยังไฟล์ต้นฉบับใน Evidence Store
    # สมมติโครงสร้าง: data_store/{tenant}/{year}/evidence/{enabler}/{filename}
    file_path = os.path.join(
        DATA_STORE_ROOT, 
        metadata.get("tenant"), 
        metadata.get("year"), 
        "evidence", 
        metadata.get("enabler").upper(), 
        filename
    )

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="ไม่พบไฟล์ต้นฉบับในระบบ")

    # 4. ส่งไฟล์กลับไปให้ UI
    return FileResponse(path=file_path, filename=filename)