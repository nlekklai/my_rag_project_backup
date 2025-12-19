# gen_seam_report.py
# Final Bug-Fixed Version - 20 ธันวาคม 2568
# แก้ไข ValueError: Unknown format code 'f' และปัญหา Import ทั้งหมด

import json
import os
import argparse
import sys
from datetime import datetime
from typing import Dict, Any, Optional, List

# Word Generation Libraries
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

# Import path_utils แบบปลอดภัย
import utils.path_utils as pu

# ==========================
# CONFIG & STYLE SAFETY
# ==========================
try:
    from config.global_vars import SEAM_ENABLER_MAP, THAI_FONT_NAME
except ImportError:
    SEAM_ENABLER_MAP = {"KM": "7.1 การจัดการความรู้ (Knowledge Management)"}
    THAI_FONT_NAME = "Angsana New"

REPORT_DATE_STR = datetime.now().strftime("%Y%m%d_%H%M%S")
DISPLAY_DATE = datetime.now().strftime("%d/%m/%Y %H:%M")

def setup_doc_style(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    style = doc.styles['Normal']
    style.font.name = THAI_FONT_NAME
    style.font.size = Pt(14)
    r = style.element.get_or_add_rPr()
    r.get_or_add_rFonts().set(qn('w:eastAsia'), THAI_FONT_NAME)

def set_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = THAI_FONT_NAME
        run.font.color.rgb = RGBColor(31, 56, 100) if level == 1 else RGBColor(50, 50, 50)

# ==========================
# REPORT SECTIONS
# ==========================

def sec_1_summary(doc, summary, enabler_name, tenant, year):
    set_h(doc, f'ส่วนที่ 1: บทสรุปเชิงยุทธศาสตร์ ({enabler_name})', 1)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    # คำนวณเปอร์เซ็นต์แบบปลอดภัย
    pct = summary.get('overall_percentage') or summary.get('percentage_achieved_run') or 0.0
    try: pct = float(pct)
    except: pct = 0.0

    data_points = [
        ("หน่วยงานที่รับการประเมิน", tenant.upper()),
        ("ปีงบประมาณ", str(year)),
        ("ระดับ Maturity (Weighted)", summary.get('Overall Maturity Level (Weighted)', 'N/A')),
        ("Performance Index (%)", f"{pct:.2f}%"),
        ("วันที่จัดทำรายงาน", DISPLAY_DATE)
    ]
    for label, val in data_points:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(val)
    doc.add_paragraph("\n")

def sec_2_gap_analysis(doc, raw_data):
    """วิเคราะห์ผลราย Level พร้อม PDCA Breakdown และ Evidence Strength"""
    set_h(doc, 'ส่วนที่ 2: วิเคราะห์ช่องว่างและคุณภาพหลักฐาน (Gap & PDCA Analysis)', 1)
    
    from collections import defaultdict
    grouped = defaultdict(list)
    for item in raw_data:
        grouped[item.get("sub_criteria_id")].append(item)

    for sid in sorted(grouped.keys()):
        levels = sorted(grouped[sid], key=lambda x: x.get('level', 0))
        set_h(doc, f"เกณฑ์ {sid}", 2)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, h in enumerate(["Level", "สถานะ", "PDCA Breakdown", "คุณภาพหลักฐาน", "เหตุผล/ข้อเสนอแนะ"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        for lv_data in levels:
            row = table.add_row().cells
            row[0].text = f"L{lv_data.get('level')}"
            row[1].text = "✅ ผ่าน" if lv_data.get('is_passed') else "❌ ไม่ผ่าน"
            
            # PDCA Breakdown
            pdca = lv_data.get('pdca_breakdown', {})
            pdca_str = f"P: {pdca.get('P','-')}\nD: {pdca.get('D','-')}\nC: {pdca.get('C','-')}\nA: {pdca.get('A','-')}"
            row[2].text = pdca_str
            
            # --- แก้ไข Bug ValueError ด้วยการแปลง Type ---
            try:
                strength = float(lv_data.get('evidence_strength', 0))
            except:
                strength = 0.0
                
            try:
                conf = float(lv_data.get('ai_confidence', 0)) * 100
            except:
                conf = 0.0
                
            row[3].text = f"Strength: {strength:.1f}/10\n(Conf: {conf:.0f}%)"
            row[4].text = lv_data.get('reason', 'ไม่พบข้อมูล')

def sec_3_roadmap(doc, sub_results):
    set_h(doc, 'ส่วนที่ 3: แผนปฏิบัติการเชิงยุทธศาสตร์ (Strategic Roadmap)', 1)
    for res in sub_results:
        sid = res.get("sub_criteria_id")
        set_h(doc, f"แนวทางยกระดับเกณฑ์ {sid}: {res.get('sub_criteria_name')}", 2)
        plans = res.get("action_plan", [])
        if not plans:
            doc.add_paragraph("สถานะ: บรรลุเป้าหมายสูงสุด (รักษามาตรฐานและนวัตกรรม)")
            continue
        for phase in plans:
            p = doc.add_paragraph()
            p.add_run(f"📌 {phase.get('phase', 'Phase')}: {phase.get('goal', 'Goal')}").bold = True
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            h = table.rows[0].cells
            h[0].text = "กิจกรรมการพัฒนา"; h[1].text = "หลักฐานยืนยัน (Outcome)"; h[2].text = "ผู้รับผิดชอบ"
            for act in phase.get('actions', []):
                for step in act.get('steps', []):
                    r = table.add_row().cells
                    r[0].text = step.get('Description', step.get('description', '-'))
                    r[1].text = step.get('Verification_Outcome', step.get('verification_outcome', '-'))
                    r[2].text = step.get('Responsible', step.get('responsible', '-'))

# ==========================
# MAIN EXECUTION
# ==========================
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("filename")
    parser.add_argument("--tenant", required=True)
    parser.add_argument("--year", type=int, required=True)
    parser.add_argument("--enabler", required=True)
    args = parser.parse_args()

    # ใช้ Logic ค้นหาไฟล์ภายในตัว (Self-Contained)
    try:
        export_root = pu.get_tenant_year_export_root(args.tenant, args.year)
        enabler_dir = os.path.join(export_root, pu._n(args.enabler))
        json_path = os.path.join(enabler_dir, args.filename)
        
        if not os.path.exists(json_path):
            if os.path.exists(enabler_dir):
                matches = [f for f in os.listdir(enabler_dir) if args.filename in f]
                if matches: json_path = os.path.join(enabler_dir, matches[0])
                else: print(f"❌ File not found in {enabler_dir}"); sys.exit(1)
            else: print(f"❌ Folder not found: {enabler_dir}"); sys.exit(1)

        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"❌ Error loading file: {e}"); sys.exit(1)

    doc = Document()
    setup_doc_style(doc)
    
    enabler_full = SEAM_ENABLER_MAP.get(args.enabler.upper(), args.enabler.upper())
    raw_data = data.get("raw_llm_results", []) or data.get("raw_results_ref", [])
    
    sec_1_summary(doc, data['summary'], enabler_full, args.tenant, args.year)
    doc.add_page_break()
    sec_2_gap_analysis(doc, raw_data)
    doc.add_page_break()
    sec_3_roadmap(doc, data['sub_criteria_results'])

    # บันทึกรายงาน
    report_root = pu.get_tenant_year_report_root(args.tenant, args.year, args.enabler)
    final_name = f"{args.tenant.upper()}_{args.year}_{args.enabler.upper()}_Strategic_Report_{REPORT_DATE_STR}.docx"
    output_path = os.path.join(report_root, final_name)
    
    doc.save(output_path)
    print(f"\n✅ รายงานสำเร็จ! (Bug Fixed)\n📍 บันทึกที่: {output_path}")

if __name__ == "__main__":
    main()