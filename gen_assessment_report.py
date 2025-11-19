import json
import os
import argparse
from typing import Dict, Any, Optional, List
from datetime import datetime

# Import libraries for DOCX generation
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# ==========================
# 1. CONFIGURATION & GLOBAL VARS
# ==========================
EXPORT_DIR = "reports"
REPORT_DATE = datetime.now().strftime("%Y-%m-%d")
THAI_FONT_NAME = "Angsana New" 

# Required Import: พยายาม Import SEAM_ENABLER_MAP จาก config/global_vars.py
try:
    from config.global_vars import SEAM_ENABLER_MAP
except ImportError:
    # Fallback/Placeholder: หากรันโค้ดนอกโครงสร้างโปรเจกต์
    print("⚠️ ไม่พบ config.global_vars. ใช้ SEAM_ENABLER_MAP จำลอง.")
    SEAM_ENABLER_MAP = {
        "KM": "7.1 การจัดการความรู้ (Knowledge Management)",
        "IT": "7.2 เทคโนโลยีดิจิทัล",
        "HR": "1.1 การบริหารทรัพยากรบุคคล",
        "GENERIC": "ตัวขับเคลื่อนทั่วไป"
    }

# ==========================
# 2. DATA LOADING & UTILITY
# ==========================

def load_data(file_path: str, file_type: str) -> Optional[Dict[str, Any]]:
    """โหลดข้อมูลจากไฟล์ JSON และจัดการข้อผิดพลาด"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        print(f"✅ โหลดไฟล์ {file_type} สำเร็จ: {file_path}")
        return data
    except FileNotFoundError:
        print(f"❌ ข้อผิดพลาดในการโหลดไฟล์ {file_type}: ไม่พบไฟล์ '{file_path}'") 
        return None
    except Exception as e:
        print(f"❌ ข้อผิดพลาดในการโหลดไฟล์ {file_path} '{file_path}': {e}") 
        return None

def setup_output_folder(file_path):
    """ตรวจสอบและสร้าง folder output"""
    output_dir = os.path.dirname(file_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

def setup_document(doc):
    """ตั้งค่าการจัดรูปแบบเอกสารโดยรวม เช่น ขอบกระดาษและฟอนต์เริ่มต้น"""
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.75) 
    section.right_margin = Inches(0.75)

    doc.styles['Normal'].font.name = THAI_FONT_NAME
    
def add_paragraph(doc, text, bold=False, italic=False, color=None, style=None):
    """ฟังก์ชัน Utility สำหรับการเพิ่ม Paragraph อย่างง่าย (พร้อมกำหนด Angsana New)"""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(text)
    
    run.font.name = THAI_FONT_NAME 
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt(11)
    return p

def set_heading(doc, text, level=1):
    """ฟังก์ชัน Utility สำหรับการเพิ่ม Heading (พร้อมกำหนด Angsana New)"""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    
    for run in p.runs:
        run.font.name = THAI_FONT_NAME 

# ==========================
# 3. REPORT GENERATION FUNCTIONS (DOCX) - Comprehensive Report
# ==========================

def generate_overall_summary_docx(document: Document, summary: Dict[str, Any], enabler_name_full: str): 
    """สร้างส่วนสรุปผลการประเมินโดยรวม (Overall) [SECTION 1] ใน DOCX"""
    overall = summary.get("summary", {})
    
    set_heading(document, f'[SECTION 1] สรุปผลการประเมิน {enabler_name_full} โดยรวม', level=1)
    
    # ดึงค่าหลักจาก Summary
    maturity_score = overall.get('Overall Maturity Score (Avg.)', 0.0)
    target_level = overall.get('target_level', 0)
    progress_percent = overall.get('percentage_achieved_run', 0.0)
    total_score = overall.get('Total Weighted Score Achieved', 0.0)
    total_possible = overall.get('Total Possible Weight', 0.0)
    
    table = document.add_table(rows=5, cols=2) 
    table.style = 'Table Grid'
    
    def add_summary_row(row_index, label, value, alignment='RIGHT'):
        table.cell(row_index, 0).text = label
        table.cell(row_index, 1).text = value
        table.cell(row_index, 0).paragraphs[0].runs[0].font.bold = True
        table.cell(row_index, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if alignment == 'RIGHT' else WD_ALIGN_PARAGRAPH.LEFT
    
    add_summary_row(0, "ตัวขับเคลื่อน (Enabler):", f"{overall.get('enabler', '-')} ({enabler_name_full})", 'LEFT') 
    add_summary_row(1, "ระดับวุฒิภาวะโดยรวม (Maturity Score):", f"{maturity_score:.2f}")
    add_summary_row(2, "ระดับเป้าหมายที่กำหนด (Target Level):", f"L{target_level}")
    add_summary_row(3, "เปอร์เซ็นต์ความคืบหน้าโดยรวม:", f"{progress_percent:.2f}%")
    add_summary_row(4, "คะแนนรวมถ่วงน้ำหนักที่ได้:", f"{total_score:.2f} / {total_possible:.2f}")
    
    # กำหนดฟอนต์ให้กับ Table Headers และ Content
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = THAI_FONT_NAME
                    run.font.size = Pt(11)

    document.add_paragraph() 

def generate_sub_criteria_status_docx(document: Document, summary: Dict[str, Any]):
    """สร้างตารางสถานะการประเมินรายเกณฑ์ย่อย [SECTION 2]"""
    sub_results = summary.get("sub_criteria_results", [])
    overall = summary.get("summary", {})
    target_level = overall.get("target_level", 0)

    document.add_heading(f'[SECTION 2] สถานะการบรรลุเป้าหมาย ({target_level}) รายเกณฑ์ย่อย', level=1)
    
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    headers = ["ID", "ชื่อเกณฑ์ย่อย", "น้ำหนัก", "Level สูงสุดที่ผ่าน", "สถานะเป้าหมาย (L{})".format(target_level), "คะแนนที่ได้"]
    for i, h in enumerate(headers):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.name = THAI_FONT_NAME 
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for sub in sub_results:
        row_cells = table.add_row().cells
        
        sub_id = sub.get("sub_criteria_id", "N/A")
        name = sub.get('sub_criteria_name', 'N/A')
        weight = sub.get('weight', 0)
        level = sub.get('highest_full_level', 0)
        score = sub.get('weighted_score', 0.0)
        
        target_achieved = sub.get('target_level_achieved', False)
        status_text = "✅ บรรลุเป้าหมาย" if target_achieved else "❌ มีช่องว่าง"
        status_color = (0x00, 0x80, 0x00) if target_achieved else (0xFF, 0x00, 0x00)
        
        row_cells[0].text = sub_id
        row_cells[1].text = name
        row_cells[2].text = str(weight)
        row_cells[3].text = f"L{level}"
        
        # สถานะเป้าหมาย (คอลัมน์ 4)
        status_cell = row_cells[4]
        status_cell.text = status_text
        status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(*status_color)
        status_cell.paragraphs[0].runs[0].font.bold = True
        
        row_cells[5].text = f"{score:.2f}"
        
        # กำหนดฟอนต์สำหรับแถวข้อมูล
        for row_cell in row_cells:
            for p in row_cell.paragraphs:
                for run in p.runs:
                    run.font.name = THAI_FONT_NAME
                    run.font.size = Pt(11)

        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph() 

def generate_action_plan_report_docx(document: Document, summary_data: Dict[str, Any]):
    """
    สร้างรายงานรายละเอียดแผนปฏิบัติการ (Action Plan) ตามแนวทาง PDCA [SECTION 3] 
    ปรับปรุงการแสดงสถานะ L1-L5 ให้ชัดเจนขึ้น
    """
    
    all_sub_criteria_results = summary_data.get("sub_criteria_results", [])
    target_level = summary_data.get("summary", {}).get("target_level", 0)

    # 1. Grouping
    gap_criteria_list = [
        sub for sub in all_sub_criteria_results if not sub.get('target_level_achieved', False)
    ]
    achieved_criteria_list = [
        sub for sub in all_sub_criteria_results if sub.get('target_level_achieved', False)
    ]
    
    # 2. Section Title
    document.add_heading('[SECTION 3] แผนปฏิบัติการเพื่อปิดช่องว่างและข้อเสนอแนะ (PDCA Approach)', level=1)
    
    # --- 3. Sub-Section: Achieved Criteria (Maintain/Sustain) ---
    document.add_heading("3.1 เกณฑ์ที่บรรลุเป้าหมาย (Good Performance & Maintain)", level=2)
    
    if achieved_criteria_list:
        add_paragraph(document, f"เกณฑ์ต่อไปนี้ได้บรรลุหรือผ่านระดับเป้าหมายที่ L{target_level} แล้ว ข้อเสนอแนะมุ่งเน้นที่การรักษาระดับวุฒิภาวะให้ยั่งยืน (Sustain) และการก้าวไปสู่ระดับที่สูงขึ้น (L{target_level+1} เป็นต้นไป).", italic=True)
        
        for sub_info in achieved_criteria_list:
            sub_id = sub_info.get("sub_criteria_id", "N/A")
            sub_name = sub_info.get('sub_criteria_name', 'N/A')
            achieved_level = sub_info.get('highest_full_level', 0)
            
            document.add_heading(f"• เกณฑ์ย่อย {sub_id}: {sub_name}", level=3) 
            add_paragraph(document, f"🎯 **สถานะ:** บรรลุเป้าหมายที่ L{target_level} (ผ่านถึง L{achieved_level})", bold=True, color=(0x00, 0x80, 0x00))

            # ACT/CHECK Component for Sustaining
            add_paragraph(document, f"✅ ACTION FOCUS: การรักษาระดับ L{achieved_level} (Sustain)", bold=True, color=(0x00, 0x80, 0x00))
            
            recommendation_table = document.add_table(rows=1, cols=3, style='Table Grid')
            header_cells = recommendation_table.rows[0].cells
            headers = ["สถานะที่ผ่าน", "ข้อเสนอแนะหลัก (Maintain)", "แนวทางปฏิบัติ"]
            for cell, h in zip(header_cells, headers):
                 cell.text = h
                 cell.paragraphs[0].runs[0].font.bold = True
                 cell.paragraphs[0].runs[0].font.name = THAI_FONT_NAME 
                 cell.paragraphs[0].runs[0].font.size = Pt(10.5)

            row_cells = recommendation_table.add_row().cells
            row_cells[0].text = f"L{target_level} ถึง L{achieved_level}"
            row_cells[1].text = "รักษาความต่อเนื่องของการดำเนินงาน และพิจารณาขยายผลเพื่อเพิ่มระดับ"
            row_cells[2].text = "ทบทวนเอกสารและหลักฐานปัจจุบันอย่างสม่ำเสมอเพื่อป้องกันการลดลงของระดับวุฒิภาวะ (De-maturity) และเริ่มวางแผนสำหรับ L{}".format(achieved_level + 1)
            
            # Set font for table
            for row in recommendation_table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = THAI_FONT_NAME
                            run.font.size = Pt(10.5)
            
            document.add_paragraph() 
    else:
        add_paragraph(document, "ℹ️ ไม่มีเกณฑ์ย่อยใดที่บรรลุเป้าหมาย L{} ในรอบการประเมินนี้".format(target_level), italic=True)
        document.add_paragraph()
        
    # --- 4. Sub-Section: Gap Criteria (Improvement) ---
    document.add_heading("3.2 เกณฑ์ที่มีช่องว่าง (Gap Closure & Improvement)", level=2)

    if not gap_criteria_list:
        add_paragraph(document, "✅ ทุกเกณฑ์ย่อยบรรลุเป้าหมายที่กำหนดแล้ว ไม่จำเป็นต้องมี Action Plan เพื่อปิดช่องว่าง", bold=True, color=(0x00, 0x80, 0x00))
    
    if gap_criteria_list: 
        add_paragraph(document, "ข้อเสนอแนะต่อไปนี้มุ่งเน้นที่การปิดช่องว่าง (Gap Closure) ระหว่างระดับวุฒิภาวะที่ผ่านสูงสุดกับระดับเป้าหมายที่กำหนด", italic=True)
        document.add_paragraph()

        for sub_info in gap_criteria_list:
            sub_id = sub_info.get("sub_criteria_id", "N/A")
            sub_name = sub_info.get('sub_criteria_name', 'N/A')
            current_level = sub_info.get('highest_full_level', 0)
            
            document.add_heading(f"• เกณฑ์ย่อย {sub_id}: {sub_name}", level=3)
            add_paragraph(document, 
                          f"🛑 **สถานะ:** บรรลุถึง L{current_level} | **ช่องว่าง:** ต้องปิดช่องว่างเพื่อบรรลุ L{target_level} (หรือสูงกว่า)", 
                          bold=True, 
                          color=(0x80, 0x00, 0x00))
            document.add_paragraph()
            
            action_plans = sub_info.get("action_plan", [])
            
            if action_plans:
                
                for i, plan_phase in enumerate(action_plans, 1):
                    
                    phase = plan_phase.get('Phase', 'N/A')
                    goal = plan_phase.get('Goal', 'N/A')
                    actions_list = plan_phase.get('Actions', [])
                    
                    add_paragraph(document, f"--- [Phase {i}] ({phase}) ---", bold=True, color=(0x44, 0x72, 0xC4))
                    
                    # Plan Component (Goal)
                    add_paragraph(document, f"🎯 PLAN: เป้าหมายหลักของ Phase นี้ (Goal)", bold=True)
                    add_paragraph(document, f"   - {goal}", style='List Bullet')

                    if actions_list:
                        
                        # Plan Component (Actions) - Changed label to DO
                        add_paragraph(document, f"💡 DO: แผนปฏิบัติการ (Action Plan) {len(actions_list)} รายการ:", bold=True) 
                        
                        action_table = document.add_table(rows=1, cols=4, style='Table Grid')
                        header_cells = action_table.rows[0].cells
                        headers = ["Level ที่ต้องบรรลุ", "ข้อเสนอแนะ (Recommendation)", "หลักฐานเป้าหมาย (Target Evidence)", "ตัวชี้วัดสำคัญ (Key Metric)"]
                        for cell, h in zip(header_cells, headers):
                             cell.text = h
                             cell.paragraphs[0].runs[0].font.bold = True
                             cell.paragraphs[0].runs[0].font.name = THAI_FONT_NAME 
                             cell.paragraphs[0].runs[0].font.size = Pt(10.5)

                        for action in actions_list:
                            row_cells = action_table.add_row().cells
                            
                            # Changed from Failed_Level to Level ที่ต้องบรรลุ
                            failed_level = action.get('Failed_Level', '-') 
                            recommendation = action.get('Recommendation', '-')
                            evidence_type = action.get('Target_Evidence_Type', '-')
                            key_metric = action.get('Key_Metric', '-')
                            
                            row_cells[0].text = f"L{failed_level}"
                            row_cells[1].text = recommendation
                            row_cells[2].text = evidence_type
                            row_cells[3].text = key_metric
                            
                            # กำหนดฟอนต์สำหรับแถวข้อมูล
                            for row_cell in row_cells:
                                for p in row_cell.paragraphs:
                                    for run in p.runs:
                                        run.font.name = THAI_FONT_NAME
                                        run.font.size = Pt(10.5)
                    
                    document.add_paragraph() 
            else:
                add_paragraph(document, ">>> [ข้อมูล]: ไม่พบ Action Plan ถูกกำหนดไว้ในส่วนนี้ โปรดกำหนดแผนเพื่อปิดช่องว่าง", style='List Bullet', color=(0x80, 0x80, 0x80))


# ==========================
# 4. RAW DETAILS REPORT FUNCTION - สร้างไฟล์แยกสำหรับ Section 4
# ==========================

def generate_raw_details_report_docx(document: Document, raw_data: Optional[Dict[str, Any]], enabler_name_full: str): 
    """สร้างรายงานการตรวจสอบความถูกต้องเชิงลึก (Raw Details) [SECTION 4] ใน DOCX"""
    
    document.add_heading('[SECTION 4] รายงานการตรวจสอบความถูกต้องเชิงลึก (Raw Details)', level=1)

    if raw_data is None:
        document.add_paragraph(f"⚠️ ไม่สามารถโหลดไฟล์ Raw Details ได้ หรือไฟล์ว่างเปล่า") 
        return
        
    assessment_details = {}
    
    if isinstance(raw_data, dict):
        # Case 1: Raw Data เป็น New combined structure (มาจาก summary_file)
        if 'sub_criteria_results' in raw_data:
             temp_dict = {}
             for sub_result in raw_data['sub_criteria_results']:
                 sub_id = sub_result.get('sub_criteria_id')
                 statements = sub_result.get('statements', [])
                 if sub_id and statements:
                      temp_dict[sub_id] = statements
             assessment_details = temp_dict
        # Case 2: Raw Data เป็น Old Raw Structure
        elif 'Assessment_Details' in raw_data:
             assessment_details = raw_data.get('Assessment_Details', {})
             
    elif isinstance(raw_data, list):
        # Case 3: Raw Data เป็น List ของ statements
        statements_list = raw_data
        temp_dict = {}
        for stmt in statements_list:
            sid = stmt.get('sub_criteria_id', 'N/A')
            if sid != 'N/A':
                if sid not in temp_dict:
                    temp_dict[sid] = []
                temp_dict[sid].append(stmt)
        assessment_details = temp_dict
        
    else:
        add_paragraph(document, f"⚠️ โครงสร้างข้อมูล Raw Details ไม่ถูกต้อง (ไม่ใช่ Dict หรือ List)") 
        return

    if not assessment_details:
         add_paragraph(document, "ℹ️ ข้อมูล Raw Details ว่างเปล่าหลังจากตรวจสอบโครงสร้าง")
         return
    
    
    for sub_id, statements in assessment_details.items():
        document.add_heading(f"รายละเอียดการประเมินเกณฑ์ย่อย: {sub_id}", level=2)
        
        table = document.add_table(rows=1, cols=6, style='Table Grid')
        # ... (Table setup for Raw Details, same as before) ...
        header_cells = table.rows[0].cells
        
        headers = ["Statement ID (Level)", "ผลการประเมิน", "Statement / Standard", "เหตุผล/วิเคราะห์", "แหล่งที่มา", "หลักฐาน/บริบท (Snippet)"] 
        for i, h in enumerate(headers):
            header_cells[i].text = h
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.name = THAI_FONT_NAME 
            header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        for statement in statements:
            if not isinstance(statement, dict):
                print(f"⚠️ ข้าม Statement ที่ไม่ใช่ Dict ใน SubID {sub_id}: {statement}")
                continue 
            
            status = "✅ PASS" if statement.get('is_pass', statement.get('pass_status', False)) else "❌ FAIL"
            level = statement.get('level', '-')
            
            reason_text = statement.get('reason', 'N/A')
            sources_list = statement.get('retrieved_sources_list', [])
            sources_text = "\n".join([
                f"{src.get('source_name', 'N/A')} (p.{src.get('location', 'N/A')})"
                for src in sources_list if isinstance(src, dict)
            ]) if sources_list else 'ไม่มีแหล่งที่มา'
            
            statement_text = statement.get('statement', 'N/A') 
            standard_text = statement.get('standard', 'N/A')   
            
            MAX_LEN_STANDARD = 150 
            # Note: We display the full statement and standard in the table cell
            
            row_cells = table.add_row().cells
            
            row_cells[0].text = f"{statement.get('statement_id', '-')}\n(L{level})"
            row_cells[1].text = status
            
            # Column 3 (Combined Statement/Standard)
            p = row_cells[2].paragraphs[0]
            
            run1 = p.add_run(statement_text)
            run1.font.name = THAI_FONT_NAME
            run1.font.size = Pt(11)
            run1.bold = True
            
            run2 = p.add_run(" / ")
            run2.font.name = THAI_FONT_NAME
            run2.font.size = Pt(11)

            run3 = p.add_run(standard_text)
            run3.font.name = THAI_FONT_NAME
            run3.font.size = Pt(11)
            run3.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Red color
            
            # Other columns
            row_cells[3].text = reason_text 
            row_cells[4].text = sources_text 
            row_cells[5].text = statement.get('context_retrieved_snippet', 'ไม่มีหลักฐานสนับสนุนที่ชัดเจน')

            # Set font for data rows (skipping column 2 which uses runs)
            for i in [0, 1, 3, 4, 5]: 
                cell = row_cells[i]
                for p_cell in cell.paragraphs:
                    for run in p_cell.runs:
                        run.font.name = THAI_FONT_NAME
                        run.font.size = Pt(11)

            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not statement.get('is_pass', statement.get('pass_status', False)):
                row_cells[1].paragraphs[0].runs[0].font.bold = True 

        document.add_paragraph() 


# ==========================
# 5. MAIN EXECUTION (Revised to output 2 files)
# ==========================
def main():
    """ฟังก์ชันหลักในการสร้างรายงานทั้งหมด"""
    
    parser = argparse.ArgumentParser(description="Generate Comprehensive Assessment Reports based on New JSON Structure.")
    parser.add_argument("--mode", choices=["all", "sub"], default="all", help="all: Generate full report. sub: Generate report for a specific sub-criteria.")
    parser.add_argument("--sub", type=str, help="SubCriteria ID (e.g., 2.2) if mode=sub.")
    parser.add_argument("--summary_file", type=str, required=True, help="Path to the Strategic/Summary JSON file (New combined structure).")
    parser.add_argument("--raw_file", type=str, required=False, default=None, help="Path to the Raw Details JSON file. If omitted, it defaults to the value of --summary_file.")
    parser.add_argument("--output_path", type=str, default="reports/Comprehensive_Report", help="Output directory and base filename prefix (e.g., reports/KM_Comprehensive_Report).")
    
    args = parser.parse_args()
    
    # 2.1 หากไม่ระบุ --raw_file ให้ใช้ --summary_file แทน (รักษาการปรับปรุงครั้งที่แล้ว)
    if args.raw_file is None:
        args.raw_file = args.summary_file
        print(f"ℹ️ ไม่ได้ระบุ --raw_file ใช้ไฟล์ Summary เป็น Raw Details แทน: {args.raw_file}")
    
    # 1. จัดการ Folder Output
    output_dir = os.path.dirname(args.output_path)
    if not output_dir:
         output_dir = EXPORT_DIR 
    setup_output_folder(output_dir)
    
    # 2. โหลดไฟล์
    summary_data_core = load_data(args.summary_file, "Strategic/Summary Core Data")
    raw_data = load_data(args.raw_file, "Raw Details Data") 
    
    if not summary_data_core or not summary_data_core.get("summary"):
        print("🚨 ไม่สามารถสร้างรายงานได้เนื่องจากไฟล์ Summary Core Data ไม่พร้อมหรือไม่มีคีย์ 'summary'")
        return
        
    # --- 3. ดึง ENABLER และกำหนดค่าเริ่มต้น ---
    enabler_id = summary_data_core.get("summary", {}).get("enabler", "GENERIC").upper() 
    enabler_name_full = SEAM_ENABLER_MAP.get(enabler_id, f"Unknown Enabler ({enabler_id})")
    
    final_summary_data = summary_data_core
    final_raw_data = raw_data
    
    # 4. การจัดการชื่อไฟล์และการกรองข้อมูล
    base_prefix = os.path.basename(args.output_path)
    if not base_prefix:
        base_prefix = f"{enabler_id}_Report"

    # 5. การกรองข้อมูลสำหรับโหมด 'sub'
    if args.mode == "sub" and args.sub:
        sub_id = args.sub.upper()
        print(f"🔹 โหมด: รายงานเฉพาะเกณฑ์ย่อย {sub_id} สำหรับ {enabler_name_full}")
        
        sub_results_list = final_summary_data.get("sub_criteria_results", [])
        
        # กรอง Summary Data
        filtered_sub_result = [s for s in sub_results_list if s.get("sub_criteria_id", "").upper() == sub_id]
        
        if filtered_sub_result:
            
            # อัปเดตโครงสร้าง Summary Data
            final_summary_data["sub_criteria_results"] = filtered_sub_result
            
            # กรอง Raw Data
            if raw_data:
                # 1. Raw Data เป็น List ของ statements
                if isinstance(raw_data, list):
                    raw_filtered = [
                        stmt for stmt in raw_data 
                        if isinstance(stmt, dict) and stmt.get('sub_criteria_id', '').upper() == sub_id
                    ]
                    final_raw_data = raw_filtered
                # 2. Raw Data เป็น Dict 
                elif isinstance(raw_data, dict):
                    assessment_details_data = raw_data.get('Assessment_Details', {})
                    # กรณีใช้ไฟล์ Summary เป็น Raw
                    if not assessment_details_data and 'sub_criteria_results' in raw_data:
                         for sub_res in raw_data['sub_criteria_results']:
                              if sub_res.get('sub_criteria_id', '').upper() == sub_id:
                                   assessment_details_data[sub_id] = sub_res.get('statements', [])
                                   break

                    if sub_id in assessment_details_data:
                        # คงโครงสร้าง Dict ที่มีแค่ Sub นั้น 
                        final_raw_data = {
                            "Assessment_Details": {sub_id: assessment_details_data.get(sub_id, [])}
                        }
                        # ต้องจัดการกรณี final_raw_data มีแค่ statements list ตรงๆ ด้วย
                        if not final_raw_data['Assessment_Details']:
                             final_raw_data = assessment_details_data.get(sub_id, [])
                    else:
                        print(f"🚨 ไม่พบ Raw Details สำหรับเกณฑ์ย่อย {sub_id} ในไฟล์ Raw Data ที่โหลดมา")
                        final_raw_data = None
            else:
                 final_raw_data = None
            
        else:
            print(f"🚨 ไม่พบเกณฑ์ย่อย {sub_id} ในไฟล์ Summary Core Data ที่โหลดมา")
            return
            
    # 6. กำหนดชื่อไฟล์ Output
    comprehensive_path = os.path.join(output_dir, f"{base_prefix}_Comprehensive_Report_{REPORT_DATE}.docx")
    detail_path = os.path.join(output_dir, f"{base_prefix}_Raw_Details_Report_{REPORT_DATE}.docx")
    
    print("-" * 50)
    print(f"🎯 ENABLER: {enabler_name_full}")
    print(f"📝 DOCX Output 1 (Comprehensive): {comprehensive_path}")
    print(f"📝 DOCX Output 2 (Raw Details): {detail_path}")
    print("-" * 50)

    # 7. สร้าง DOCX Report 1: Comprehensive (Sections 1, 2, 3)
    try:
        document = Document()
        setup_document(document)
        
        # --- Header ---
        set_heading(document, f"รายงานผลการประเมิน {enabler_name_full}", level=1)
        add_paragraph(document, f"วันที่สร้างรายงาน: {REPORT_DATE}", style='Caption')
        document.add_paragraph()
        
        # Section 1: Overall Summary
        generate_overall_summary_docx(document, final_summary_data, enabler_name_full)
        
        # Section 2: Sub-Criteria Status 
        generate_sub_criteria_status_docx(document, final_summary_data) 
        
        # Section 3: Action Plan & PDCA Recommendations
        generate_action_plan_report_docx(document, final_summary_data)
        
        document.save(comprehensive_path)
        print(f"✅ สร้างรายงาน DOCX [Comprehensive] สำเร็จ: {comprehensive_path}")
    except Exception as e:
        print(f"❌ ข้อผิดพลาดในการสร้าง DOCX [Comprehensive] Report: {e}")
        
    # 8. สร้าง DOCX Report 2: Raw Details (Section 4)
    try:
        detail_doc = Document()
        setup_document(detail_doc)
        
        # --- Header ---
        set_heading(detail_doc, f"รายงานการตรวจสอบความถูกต้องเชิงลึก {enabler_name_full}", level=1)
        add_paragraph(detail_doc, f"วันที่สร้างรายงาน: {REPORT_DATE}", style='Caption')
        detail_doc.add_paragraph()
        
        # Section 4: Raw Details
        generate_raw_details_report_docx(detail_doc, final_raw_data, enabler_name_full)

        detail_doc.save(detail_path)
        print(f"✅ สร้างรายงาน DOCX [Raw Details] สำเร็จ: {detail_path}")
    except Exception as e:
        print(f"❌ ข้อผิดพลาดในการสร้าง DOCX [Raw Details] Report: {e}")

if __name__ == "__main__":
    main()