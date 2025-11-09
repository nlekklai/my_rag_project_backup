#gen_gpt_report.py
import json
import os
import argparse
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ==========================
# CONFIG
# ==========================
EXPORT_DIR = "reports"
REPORT_DATE = datetime.now().strftime("%Y-%m-%d")
THAI_FONT_NAME = "Angsana New" # กำหนดฟอนต์ภาษาไทย

# ==========================
# UTILITIES
# ==========================
def setup_document(doc):
    """Sets up document-wide formatting like margins and default font."""
    # 1. ตั้งค่า Margins: ลดระยะขอบซ้าย/ขวาเหลือ 0.75 นิ้ว
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.75) 
    section.right_margin = Inches(0.75)

    # 2. ตั้งค่า Default Font เป็น Angsana New
    # 💡 FIX: ลบเงื่อนไข if ที่ทำให้เกิด TypeError และกำหนดค่า font name โดยตรง
    # การกำหนดค่าโดยตรงจะบังคับให้ใช้ Angsana New ซึ่งเป็นความต้องการของเรา
    doc.styles['Normal'].font.name = THAI_FONT_NAME
        
def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    # กำหนดฟอนต์ให้กับ Heading
    for run in p.runs:
        run.font.name = THAI_FONT_NAME
    return p

def add_paragraph(doc, text, bold=False, italic=False, color=None, style=None): 
    """ฟังก์ชัน Utility สำหรับการเพิ่ม Paragraph พร้อมการตั้งค่าสไตล์"""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(text)
    
    # กำหนดฟอนต์ให้กับ Run
    run.font.name = THAI_FONT_NAME 
    
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt(11)
    return p

def save_doc(doc, name):
    os.makedirs(EXPORT_DIR, exist_ok=True)
    output_path = os.path.join(EXPORT_DIR, name)
    doc.save(output_path)
    print(f"✅ Created: {output_path}")

def load_data(file_path: str) -> Dict[str, Any]:
    """โหลดไฟล์ JSON พร้อมจัดการข้อผิดพลาด"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            print(f"✅ โหลดไฟล์: {file_path}")
            return json.load(f) 
    except FileNotFoundError:
        print(f"❌ Error: ไม่พบไฟล์อินพุต '{file_path}'")
        return {} 
    except Exception as e:
        print(f"❌ Error loading {file_path}: {e}")
        return {}

def flatten_raw_data(raw_data_dict: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    ดึง Statement ทั้งหมดออกมาจาก Raw Data Dictionary ให้อยู่ในรูปแบบ List ที่สามารถวนซ้ำได้ง่าย
    (รองรับทั้งโครงสร้าง Dict และ List)
    """
    statements = []
    
    if isinstance(raw_data_dict, dict):
        details = raw_data_dict.get("Assessment_Details")
        if isinstance(details, dict):
            for sub_id_statements in details.values():
                if isinstance(sub_id_statements, list):
                    statements.extend(sub_id_statements)
    
    elif isinstance(raw_data_dict, list):
        statements = raw_data_dict
        
    return statements

# ==========================
# REPORT 1: ภาพรวม (Overall)
# ==========================
def create_overall_report(summary):
    if not summary: return
    doc = Document()
    setup_document(doc) # <--- เรียกใช้ setup_document
    set_heading(doc, "รายงานภาพรวมการจัดการความรู้ (KM)")

    overall = summary.get("Overall", {})
    add_paragraph(doc, f"วันที่จัดทำรายงาน: {REPORT_DATE}")
    add_paragraph(doc, f"คะแนนรวมถ่วงน้ำหนัก: {overall.get('total_weighted_score', '-'):.2f}")
    add_paragraph(doc, f"คะแนนเต็มรวม: {overall.get('total_possible_weight', '-'):.2f}")
    add_paragraph(doc, f"ร้อยละความก้าวหน้า: {float(overall.get('overall_progress_percent', 0)):.2f}%")
    add_paragraph(doc, f"ระดับความเป็นผู้ใหญ่ (Maturity Score): {overall.get('overall_maturity_score', '-'):.2f}")

    add_paragraph(doc, "\nรายละเอียดเกณฑ์ย่อย (Sub-Criteria):", style="List Bullet") 
    for sid, sdata in summary.get("SubCriteria_Breakdown", {}).items():
        add_paragraph(doc, f"{sid} - {sdata.get('name', sdata.get('topic', 'N/A'))}", bold=True)
        add_paragraph(doc, f"   คะแนน: {sdata.get('score', 0):.2f}/{sdata.get('weight', 0):.2f}")
        add_paragraph(doc, f"   ระดับสูงสุดที่ผ่านเต็ม: L{sdata.get('highest_full_level', 0)}")
        add_paragraph(doc, f"   ช่องว่างพัฒนา: {'มี ❌' if sdata.get('development_gap') else 'ไม่มี ✅'}")
    save_doc(doc, "KM_Report_Overall.docx")

# ==========================
# REPORT 2: GAP Analysis
# ==========================
def create_gap_report(summary):
    if not summary: return
    doc = Document()
    setup_document(doc) # <--- เรียกใช้ setup_document
    set_heading(doc, "รายงานช่องว่างการพัฒนา (KM Gap Analysis)")

    breakdown = summary.get("SubCriteria_Breakdown", {})
    gaps_with_id = [(sid, s) for sid, s in breakdown.items() if s.get("development_gap")]
    
    if not gaps_with_id:
        add_paragraph(doc, "ไม่พบช่องว่างการพัฒนา ✅", bold=True)
    else:
        for sid, s in gaps_with_id:
            sub_name = s.get('name', s.get('topic', 'N/A'))
            current_level = s.get('highest_full_level', 0)
            
            # ปรับ Heading: ระบุ ID และ Level
            heading_text = f"Gap: {sid} - {sub_name} (Highest Full Level: L{current_level})"
            set_heading(doc, heading_text, level=2)

            add_paragraph(doc, f"คะแนน: {s.get('score', 0):.2f} / {s.get('weight', 0):.2f}")
            
            # เน้นย้ำ Action Item (ซึ่งคือการปิด Gap ไป Level ถัดไป)
            add_paragraph(doc, "🎯 **แผนงานเพื่อปิดช่องว่างปัจจุบัน (Action Plan)**", bold=True, color=(0x00, 0x00, 0xFF)) # Blue color for action
            # Action Item/Comment from Summary (นี่คือข้อเสนอแนะสำหรับการไป Level ถัดไป)
            add_paragraph(doc, s.get('action_item', 'ไม่มีข้อเสนอแนะสำหรับการปิดช่องว่างปัจจุบัน'))

            doc.add_paragraph("\n")

            # ข้อมูลเชิงลึก L5
            if "evidence_summary_L5" in s:
                add_paragraph(doc, "ข้อมูลเชิงลึกหลักฐานระดับ L5 (เป้าหมายสูงสุด):", bold=True, color=(0x00, 0x70, 0xC0))
                add_paragraph(doc, "สรุปหลักฐาน L5:", italic=True)
                add_paragraph(doc, s["evidence_summary_L5"].get("summary", "ไม่มีสรุป L5"))
                add_paragraph(doc, "ข้อเสนอแนะสำหรับ L5:", italic=True)
                add_paragraph(doc, s["evidence_summary_L5"].get("suggestion_for_next_level", "ไม่มีข้อเสนอแนะ"))
                doc.add_paragraph("\n")
            
            # ข้อมูลเชิงลึก L4
            if "evidence_summary_L4" in s:
                add_paragraph(doc, "ข้อมูลเชิงลึกหลักฐานระดับ L4:", bold=True, color=(0x00, 0x70, 0xC0))
                add_paragraph(doc, "สรุปหลักฐาน L4:", italic=True)
                add_paragraph(doc, s["evidence_summary_L4"].get("summary", "ไม่มีสรุป L4"))
                doc.add_paragraph("\n")

            doc.add_paragraph("-" * 50) 
    save_doc(doc, "KM_Report_Gap.docx")

# ==========================
# REPORT 3: หลักฐานรายข้อ (Evidence Detail)
# ==========================
def create_evidence_detail_report(raw_statements_list: List[Dict[str, Any]]):
    if not raw_statements_list: return
    doc = Document()
    setup_document(doc) # <--- เรียกใช้ setup_document
    set_heading(doc, "รายงานหลักฐานรายข้อ (KM Evidence Details)")

    for item in raw_statements_list:
        sid = item.get("statement_id")
        
        is_pass = item.get("is_pass", item.get("pass_status", False))
        status_text = 'ผ่าน ✅' if is_pass else 'ไม่ผ่าน ❌'
        status_color = (0x00, 0x70, 0xC0) if is_pass else (0xFF, 0x00, 0x00) 

        add_paragraph(doc, f"รหัสข้อ: {sid}", bold=True)
        add_paragraph(doc, f"หัวข้อ: {item.get('standard', item.get('statement', 'N/A'))}")
        add_paragraph(doc, f"ระดับ: {item.get('level')}")
        add_paragraph(doc, f"ผลการประเมิน: {status_text}", bold=True, color=status_color)
        add_paragraph(doc, f"เหตุผล: {item.get('reason', 'N/A')}")
        
        snippet = item.get("context_retrieved_snippet")
        if snippet:
            add_paragraph(doc, f"Snippet (หลักฐาน): {snippet}")

        if item.get("retrieved_sources_list"):
            add_paragraph(doc, "แหล่งหลักฐานที่ใช้:")
            for src in item["retrieved_sources_list"]:
                source_name = src.get('source_name', 'N/A')
                location = src.get('location', 'N/A')
                add_paragraph(doc, f" - {source_name} (p.{location})")
        doc.add_paragraph("-" * 50)
    save_doc(doc, "KM_Report_EvidenceDetails.docx")

# ==========================
# REPORT 4: Executive Summary
# ==========================
def create_executive_summary(summary):
    if not summary: return
    doc = Document()
    setup_document(doc) # <--- เรียกใช้ setup_document
    set_heading(doc, "รายงานสรุปสำหรับผู้บริหาร (Executive Summary)")

    overall = summary.get("Overall", {})
    add_paragraph(doc, f"คะแนนรวม: {overall.get('total_weighted_score', 0):.2f}/{overall.get('total_possible_weight', 0):.2f}")
    add_paragraph(doc, f"ร้อยละความสำเร็จ: {overall.get('overall_progress_percent', 0):.2f}%")
    add_paragraph(doc, f"ระดับความเป็นผู้ใหญ่: {overall.get('overall_maturity_score', 0):.2f}")
    add_paragraph(doc, "")

    breakdown = summary.get("SubCriteria_Breakdown", {})
    if breakdown:
        # Strength: Top 3 highest scoring
        add_paragraph(doc, "✅ จุดแข็งที่โดดเด่น (Top Strengths):", bold=True, color=(0x00, 0x70, 0xC0))
        top_strengths = sorted(breakdown.values(), key=lambda x: x.get("score", 0), reverse=True)[:3]
        for s in top_strengths:
            sub_name = s.get('name', s.get('topic', 'N/A')) 
            add_paragraph(doc, f"- {sub_name} ({s.get('score', 0):.2f}/{s.get('weight', 0):.2f})", style="List Bullet")

        add_paragraph(doc, "")
        
        # Weakness: Top 3 with Gap (or lowest scoring with Gap)
        add_paragraph(doc, "⚠️ จุดที่ควรพัฒนา (Development Areas):", bold=True, color=(0xFF, 0x00, 0x00))
        gaps = [s for s in breakdown.values() if s.get("development_gap")]
        top_weaknesses = sorted(gaps, key=lambda x: x.get("score", 0))[:3]
        for s in top_weaknesses:
            sub_name = s.get('name', s.get('topic', 'N/A'))
            add_paragraph(doc, f"- {sub_name} (ระดับสูงสุดผ่าน: L{s.get('highest_full_level', 0)})", style="List Bullet")

    save_doc(doc, "KM_Report_ExecutiveSummary.docx")

# ==========================
# MAIN
# ==========================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Comprehensive KM Reports")
    parser.add_argument("--summary_file", type=str, required=True, help="Path to the Summary JSON file.")
    parser.add_argument("--raw_file", type=str, required=True, help="Path to the Raw Details JSON file.")
    args = parser.parse_args()

    # 1. โหลดข้อมูล
    summary_data = load_data(args.summary_file)
    raw_data_dict = load_data(args.raw_file)
    
    # 2. จัดระเบียบ Raw Data (แก้ไข Bug)
    raw_statements_list = flatten_raw_data(raw_data_dict)

    if not summary_data or (isinstance(summary_data, dict) and not summary_data.get("Overall")):
        print("\n🚨 ไม่สามารถสร้างรายงานได้: ข้อมูล Summary ไม่พร้อม")
    else:
        # 3. สร้างรายงาน
        print("\n--- เริ่มสร้างรายงาน DOCX (พร้อมปรับ Font/Margin) ---")
        create_overall_report(summary_data)
        create_gap_report(summary_data)
        create_evidence_detail_report(raw_statements_list) 
        create_executive_summary(summary_data)

        print("\n🎉 รายงานทั้งหมดถูกสร้างเสร็จสมบูรณ์ในโฟลเดอร์ reports/ แล้วครับ! (ใช้ Angsana New และขอบแคบลง)")