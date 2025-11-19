#!/usr/bin/env python3
# gen_report.py (v31 - CAPPED Status Fix / Final Merge)
"""
สร้างรายงานผลการประเมิน SE-AM ในรูปแบบ Docx ที่สมบูรณ์แบบ

ปรับปรุงล่าสุด:
- **[CRITICAL FIX / V31]** แก้ไข Logic ใน `_add_level_status_summary_table` ให้แสดงสถานะ '⚠️ FAIL (CAPPED)' 
  อย่างถูกต้อง โดยอ้างอิงจากฟิลด์ 'is_capped' ที่มาจาก Engine เพื่อระบุความล้มเหลวจาก Dependency (ระดับต่ำกว่าไม่ผ่าน)
- **[PRESERVE / V30]** รวมตารางสรุปสถานะรายระดับ (`_add_level_status_summary_table`) และตาราง PDCA Breakdown เข้าด้วยกันเป็นตารางเดียว ตามคำขอของผู้ใช้
- **[PRESERVE / V30]** คงไว้ซึ่งการแก้ไขปัญหาค่า PDCA Score (Achieved/Required) โดยการดึงค่า Achieved Score จากฟิลด์ 'score_achieved' ในข้อมูล JSON โดยตรง

"""

import os
import sys
import json
import argparse
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

# จำเป็นต้องติดตั้งไลบรารี python-docx ก่อน: pip install python-docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("❌ ERROR: กรุณาติดตั้งไลบรารี python-docx ก่อน: pip install python-docx")
    sys.exit(1)


# -------------------------
# CONFIG / การตั้งค่า
# -------------------------
EXPORT_DIR = "reports"
REPORT_DATE = datetime.now().strftime("%Y-%m-%d_%H%M%S")
DISPLAY_DATE = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
THAI_FONT = "Angsana New"
FALLBACK_FONT = "Calibri"

# Font size
DEFAULT_FONT_SIZE = 14
TABLE_FONT_SIZE = 14

# กำหนดสีสำหรับรายงาน
COLOR_HEADER = RGBColor(0x00, 0x70, 0xC0)  # น้ำเงินเข้ม
COLOR_ACCENT = RGBColor(0xFF, 0x99, 0x00)  # ส้ม
COLOR_BAD = RGBColor(0xFF, 0x00, 0x00)     # แดง
COLOR_GOOD = RGBColor(0x00, 0x80, 0x00)    # เขียว
COLOR_NEUTRAL = RGBColor(0x33, 0x33, 0x33) # เทาเข้ม
COLOR_WARN = RGBColor(0xCC, 0x66, 0x00)    # ส้มอมน้ำตาล

SEAM_ENABLER_MAP = {
    "KM": "7.1 การจัดการความรู้ (Knowledge Management)",
    "IT": "7.2 เทคโนโลยีดิจิทัล",
    "HR": "1.1 การบริหารทรัพยากรบุคคล",
    "GENERIC": "ตัวขับเคลื่อนทั่วไป"
}
SOURCE_RE = re.compile(r'\[SOURCE:\s*(.+?)\s*\(ID:([0-9a-f]+)[^\)]*\)\s*\]', re.IGNORECASE)
SNIPPET_MAX_CHARS = 1000 

# -------------------------
# UTILITIES / ฟังก์ชันช่วยเหลือ
# -------------------------
def load_json(path: str) -> Optional[Dict[str, Any]]:
    """โหลดข้อมูล JSON จากไฟล์พร้อมจัดการข้อผิดพลาด"""
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        # print(f"❌ Error loading JSON '{path}': {e}") # Suppress for clean terminal
        return None

def ensure_output_dir(path: str):
    """ตรวจสอบและสร้างไดเรกทอรีเอาต์พุตหากยังไม่มี"""
    d = os.path.dirname(path)
    if d and not os.path.exists(d):
        os.makedirs(d, exist_ok=True)

def set_font_for_run(run, size: int = DEFAULT_FONT_SIZE, bold: bool = False, color: Optional[RGBColor] = None, font_name: str = THAI_FONT):
    """ตั้งค่าฟอนต์สำหรับ run ใน Docx"""
    run.font.size = Pt(size)
    run.font.bold = bold
    try:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        run.font.name = FALLBACK_FONT
    if color:
        run.font.color.rgb = color

def add_paragraph(doc_or_cell, text: str, size: int = DEFAULT_FONT_SIZE, bold: bool = False, color: Optional[RGBColor] = None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """เพิ่มย่อหน้าในเอกสารหรือเซลล์พร้อมกำหนดสไตล์"""
    if hasattr(doc_or_cell, "add_paragraph"):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
        if len(doc_or_cell.paragraphs) > 1 and not doc_or_cell.paragraphs[0].text.strip():
             doc_or_cell.paragraphs[0].clear()

    p.alignment = align
    run = p.add_run(text)
    set_font_for_run(run, size=size, bold=bold, color=color)
    return p

def add_heading_center(doc, text: str, level: int = 1):
    """เพิ่มหัวข้อจัดกึ่งกลางพร้อมกำหนดสไตล์"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    size = 18 if level == 1 else 16
    set_font_for_run(run, size=size, bold=True, color=COLOR_HEADER)
    return p

def safe_get(d: Dict[str, Any], *keys, default=None):
    """ดึงค่าจาก Dict ที่ซ้อนกันอย่างปลอดภัย"""
    cur = d
    for k in keys:
        if isinstance(cur, list) and isinstance(k, int) and 0 <= k < len(cur):
             cur = cur[k]
        elif isinstance(cur, dict):
            cur = cur.get(k, default)
        else:
            return default
    return cur

def extract_sources_from_stmt(stmt: Dict[str, Any]) -> List[str]:
    """ดึงรายการแหล่งที่มาจากข้อมูล Statement"""
    sources_list = []
    srcs = stmt.get('retrieved_full_source_info') or stmt.get('retrieved_full_sources') or []
    for s in srcs:
        name = s.get('source') or s.get('file_name') or s.get('title')
        page = s.get('page') or s.get('chunk_index')
        suffix = f" (p.{page})" if page else ""
        if name:
            sources_list.append(f"{name}{suffix}")

    if not sources_list:
        ctx = stmt.get('aggregated_context_used', '')
        matches = SOURCE_RE.findall(ctx)
        for name, _ in matches:
            sources_list.append(name)
    
    # ลบรายการซ้ำ
    sources_list = sorted(list(set(sources_list)))

    return sources_list

def clean_snippet_text(text: str) -> str:
    """ทำความสะอาดข้อความ Snippet ดิบเพื่อให้อ่านง่ายขึ้น (จัดการ Newline/Whitespace)"""
    if not text:
        return ""
    # 1. แทนที่ Newline/Carriage return ด้วยช่องว่าง
    text = text.replace('\n', ' ').replace('\r', ' ')
    # 2. แทนที่ช่องว่างหลายช่องด้วยช่องว่างเดียว
    text = re.sub(r'\s+', ' ', text).strip()
    return text

# -------------------------
# REPORT BUILDERS (PART A)
# -------------------------
def _add_executive_summary_table(doc: Document, summary: Dict[str, Any]):
    """เพิ่มตารางสรุป Executive Summary"""
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(2.4)
    table.columns[1].width = Inches(3.6)

    cell00, cell01 = table.cell(0, 0), table.cell(0, 1)
    cell00.merge(table.cell(2, 0))
    add_paragraph(cell00, "Executive Summary", size=16, bold=True, color=COLOR_HEADER, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    maturity_level = str(safe_get(summary, "Overall Maturity Level (Weighted)", default="N/A"))
    add_paragraph(cell01, "ระดับวุฒิภาวะองค์กรรวม", size=14, color=COLOR_NEUTRAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(cell01, maturity_level, size=24, bold=True, color=COLOR_ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell01.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell11 = table.cell(1, 1)
    add_paragraph(cell11, "คะแนนรวม (ถ่วงน้ำหนัก)", size=14, bold=True, color=COLOR_NEUTRAL)
    tws = safe_get(summary, 'Total Weighted Score Achieved', default=0)
    tpw = safe_get(summary, 'Total Possible Weight', default=0)
    add_paragraph(cell11, f"{tws:.2f}/{tpw:.2f}", size=14, bold=True, color=COLOR_HEADER)

    cell21 = table.cell(2, 1)
    add_paragraph(cell21, "ความคืบหน้ารวม", size=14, bold=True, color=COLOR_NEUTRAL)
    pct = safe_get(summary, "Overall Progress Percentage (0.0 - 1.0)", default=0.0) * 100
    add_paragraph(cell21, f"{pct:.1f}%", size=14, bold=True, color=COLOR_HEADER)

def _add_sub_criteria_summary_table(doc: Document, sub_results: List[Dict[str, Any]]):
    """
    เพิ่มตารางสรุปผลการประเมินรายเกณฑ์ย่อยทั้งหมด (สำหรับโหมด ALL)
    """
    if not sub_results or len(sub_results) <= 1:
        # ไม่ต้องแสดงตารางสรุปหากมีแค่เกณฑ์ย่อยเดียว
        return

    add_paragraph(doc, "📊 สรุปผลการประเมินรายเกณฑ์ย่อย (Sub-Criteria Summary)", size=16, bold=True, color=COLOR_HEADER)
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.allow_autofit = True

    # Column Headers
    headers = [
        "Sub-ID", 
        "เกณฑ์ย่อย", 
        "น้ำหนัก (W)", 
        "ระดับวุฒิภาวะ (L)", 
        "คะแนน (S/W)", 
        "สถานะเป้าหมาย"
    ]
    
    # Column Widths (Adjusted for better fit)
    table.columns[0].width = Inches(0.8)  # Sub-ID
    table.columns[1].width = Inches(3.0)  # Name
    table.columns[2].width = Inches(0.8)  # Weight
    table.columns[3].width = Inches(1.0)  # Maturity Level
    table.columns[4].width = Inches(1.2)  # Weighted Score
    table.columns[5].width = Inches(1.0)  # Target Status

    for i, h in enumerate(headers):
        add_paragraph(table.rows[0].cells[i], h, size=TABLE_FONT_SIZE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Populate rows
    # Sorting by sub_criteria_id ensures a logical display order (e.g., 1.1, 2.1, 2.2, ...)
    for s in sorted(sub_results, key=lambda x: x.get('sub_criteria_id', '')):
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        sub_id = s.get('sub_criteria_id', '-')
        sub_name = s.get('sub_criteria_name', '-')
        weight = s.get('weight', 0)
        highest_lvl = s.get('highest_full_level', 0)
        weighted_score = s.get('weighted_score', 0)
        target_achieved = s.get('target_level_achieved', False)

        # Sub-ID (Col 0)
        add_paragraph(cells[0], sub_id, size=TABLE_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Name (Col 1)
        add_paragraph(cells[1], sub_name, size=TABLE_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.LEFT)
        
        # Weight (Col 2)
        add_paragraph(cells[2], str(weight), size=TABLE_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Maturity Level (Col 3)
        add_paragraph(cells[3], f"L{highest_lvl}", size=TABLE_FONT_SIZE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Weighted Score (Col 4)
        score_text = f"{weighted_score:.2f}/{weight:.2f}"
        add_paragraph(cells[4], score_text, size=TABLE_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Target Achieved (Col 5)
        status_text = "✅ PASS" if target_achieved else "❌ FAIL"
        status_color = COLOR_GOOD if target_achieved else COLOR_BAD
        add_paragraph(cells[5], status_text, size=TABLE_FONT_SIZE, bold=True, color=status_color, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph("")


def build_overall_summary(doc: Document, summary: Dict[str, Any], sub_results: List[Dict[str, Any]], enabler_full: str):
    """สร้างส่วนสรุปภาพรวม (Part A)"""
    add_heading_center(doc, f"รายงานผลการประเมินวุฒิภาวะองค์กร (SE-AM) - {enabler_full}", level=1)
    add_paragraph(doc, f"วันที่สร้างรายงาน: {DISPLAY_DATE}", size=12, color=COLOR_NEUTRAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("")

    _add_executive_summary_table(doc, summary)
    doc.add_paragraph("")
    
    # Sub-Criteria Summary Table
    _add_sub_criteria_summary_table(doc, sub_results)
    
    doc.add_paragraph("")


    add_paragraph(doc, "📈 จุดแข็งเด่น (Top strengths)", size=16, bold=True, color=COLOR_HEADER)
    strengths = sorted(sub_results, key=lambda s: s.get('weighted_score', 0), reverse=True)[:3]
    if strengths:
        for s in strengths:
            text = f"• {s.get('sub_criteria_id','-')} {s.get('sub_criteria_name','-')} — L{s.get('highest_full_level',0)} — {s.get('weighted_score',0):.2f}/{s.get('weight',0):.2f}"
            add_paragraph(doc, text, size=14)
    else:
        add_paragraph(doc, "ไม่มีข้อมูลย่อสำหรับจุดแข็ง", size=14, color=COLOR_NEUTRAL)

    doc.add_paragraph("")

    add_paragraph(doc, "🚨 จุดที่ควรพัฒนา (Top gaps)", size=16, bold=True, color=COLOR_BAD)
    gaps = [s for s in sub_results if not s.get('target_level_achieved', True)]
    if gaps:
        target_level = safe_get(summary, 'target_level', default='N/A')
        for s in gaps[:5]:
            text = f"• {s.get('sub_criteria_id','-')} {s.get('sub_criteria_name','-')} — L{s.get('highest_full_level',0)} — ต้องการ L{target_level}"
            add_paragraph(doc, text, size=14)
    else:
        add_paragraph(doc, "✓ เกณฑ์ย่อยที่ถูกประเมินบรรลุเป้าหมายแล้ว", size=14, color=COLOR_GOOD)

    doc.add_page_break()


# -------------------------
# REPORT BUILDERS (PART B) - Detailed Section Helpers
# -------------------------

def _add_sub_criteria_meta(doc: Document, sub_id: str, sub_meta: Dict[str, Any], target_level_value: int):
    """
    เพิ่มหัวข้อและข้อมูลเมตาของเกณฑ์ย่อย พร้อมแสดงสถานะการบรรลุเป้าหมาย
    """
    sub_name = sub_meta.get('sub_criteria_name','(ชื่อไม่ระบุ)')
    header_text = f"=== เกณฑ์ย่อย {sub_id}: {sub_name} ==="
    add_paragraph(doc, header_text, size=16, bold=True, color=COLOR_HEADER)
    
    target_achieved = sub_meta.get('target_level_achieved', False)
    
    # FIX: Ensure target_level_value is an integer for safe calculation
    try:
        target_lvl_int = int(target_level_value)
    except (TypeError, ValueError):
        target_lvl_int = 0

    target_lvl_display = f"L{target_lvl_int}" if target_lvl_int else 'N/A'
    highest_lvl = sub_meta.get('highest_full_level',0)
    weighted_score = sub_meta.get('weighted_score',0)
    weight = sub_meta.get('weight',0)

    # --- Narrative Summary / คำบรรยายสรุป ---
    if target_achieved:
        narrative = (
            f"เกณฑ์ย่อยนี้แสดงผลลัพธ์ที่ยอดเยี่ยม โดยสามารถบรรลุเป้าหมายที่ตั้งไว้ที่ **{target_lvl_display}** "
            f"ด้วยการดำเนินงานที่แข็งแกร่งจนถึง **ระดับวุฒิภาวะสูงสุดที่ L{highest_lvl}** "
            f"โดยได้รับ **คะแนนถ่วงน้ำหนักเต็มจำนวน {weighted_score:.2f} จาก {weight:.2f}** คะแนน "
            f"ซึ่งบ่งชี้ว่าหลักฐานการดำเนินงานครอบคลุมครบถ้วนตามวงจร PDCA ในทุกระดับ"
        )
        status_text = f"✅ สถานะ: บรรลุเป้าหมาย (L{highest_lvl} ≥ {target_lvl_display})"
        status_color = COLOR_GOOD
    else:
        # ใช้ค่า target_lvl_int ที่เป็น integer สำหรับการคำนวณ
        gap_lvl = target_lvl_int - highest_lvl 
        
        narrative = (
            f"เกณฑ์ย่อยนี้ **ยังไม่บรรลุเป้าหมาย** ที่ตั้งไว้ที่ **{target_lvl_display}** "
            f"โดยระดับวุฒิภาวะสูงสุดที่ทำได้ในปัจจุบันคือ **L{highest_lvl}** "
            f"ได้รับคะแนนถ่วงน้ำหนักเพียง **{weighted_score:.2f} จาก {weight:.2f}** คะแนน "
            f"องค์กรจำเป็นต้องพัฒนาหลักฐานเพิ่มเติมอีก **{max(1, gap_lvl)} ระดับ** (จาก L{highest_lvl+1}) เพื่อให้บรรลุเป้าหมาย"
        )
        status_text = f"❌ สถานะ: ยังไม่บรรลุเป้าหมาย (ต้องการ L{target_lvl_display})"
        status_color = COLOR_BAD
        
    add_paragraph(doc, narrative, size=14, color=COLOR_NEUTRAL)
    doc.add_paragraph("")
    
    # --- Existing: Key Metrics (Keep it concise) ---
    add_paragraph(doc, status_text, size=14, bold=True, color=status_color)
    meta_text = f"Highest Full Level: L{highest_lvl} | Target Level: {target_lvl_display} | Weighted Score: {weighted_score:.2f}/{weight:.2f}"
    add_paragraph(doc, meta_text, size=12, color=COLOR_NEUTRAL)
    doc.add_paragraph("")

def _get_required_score(level: int) -> int:
    """Helper function to determine the required score (R) for PDCA (Achieved/Required)"""
    # NOTE: ฟังก์ชันนี้ถูกแทนที่ใน v30 ด้วยการดึงค่าจาก item.get('pdca_score_required', 8)
    # แต่ยังคงเก็บไว้เพื่อเป็น fallback หากข้อมูล JSON ไม่สมบูรณ์
    if level == 1: return 1
    if level == 2: return 2
    if level == 3: return 4
    if level == 4: return 6
    if level == 5: return 8
    return 0


def _add_level_status_summary_table(doc: Document, sub_meta: Dict[str, Any]):
    """
    V31 FIX: สร้างตารางสรุปผลการประเมินรายระดับ (L1-L5) รวม PDCA (A/R), P, D, C, A, CAPPED 
    และเหตุผลโดยสรุป โดยใช้ฟิลด์ 'is_capped' ที่มาจาก Engine โดยตรง
    (รวมตารางสรุปสถานะรายระดับและ PDCA Breakdown เข้าด้วยกัน)
    """
    sub_id = sub_meta.get('sub_criteria_id', 'N/A')
    highest_lvl = sub_meta.get('highest_full_level', 0)
    raw_results_ref = sub_meta.get('raw_results_ref', []) 

    evaluated_levels = [
        res for res in raw_results_ref 
        if res.get('level') is not None and res.get('level') > 0
    ]

    if not evaluated_levels:
        add_paragraph(doc, "ไม่พบผลการประเมินระดับย่อยใด ๆ สำหรับเกณฑ์ย่อยนี้", size=TABLE_FONT_SIZE, color=COLOR_NEUTRAL)
        doc.add_paragraph("")
        return

    add_paragraph(doc, f"📋 สรุปสถานะรายระดับ พร้อม PDCA Score (Achieved/Required), Breakdown และเหตุผล: {sub_id}", size=14, bold=True, color=COLOR_HEADER)
    
    # สร้างตาราง 9 คอลัมน์ (L, Status, A/R, P, D, C, A, CAPPED, Reason)
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    table.allow_autofit = True
    
    # Headers
    headers = [
        "ระดับ (L)", 
        "สถานะ (Result)", 
        "PDCA (A/R)", 
        "P", "D", "C", "A", # PDCA Breakdown
        "CAPPED", 
        "เหตุผลโดยสรุปจากระบบ"
    ] 
    # กำหนดความกว้างคอลัมน์ (รวมประมาณ 7.1 นิ้ว)
    widths = [Inches(0.6), Inches(0.8), Inches(1.0), Inches(0.5), Inches(0.5), Inches(0.5), Inches(0.5), Inches(0.7), Inches(2.5)] 
    
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        add_paragraph(cell, h, size=TABLE_FONT_SIZE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        try:
            cell.width = widths[i]
        except:
             pass

    # Populate rows
    for item in evaluated_levels:
        cells = table.add_row().cells
        
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        level = item.get('level', 0)
        pdca = item.get('pdca_breakdown', {})
        
        # 1. Status Derivation and Reason (FIXED: ใช้สถานะ is_capped ที่มาจาก Engine โดยตรง)
        
        # 🟢 FIX 1: ดึงสถานะ is_capped ที่คำนวณจาก Engine โดยตรง
        is_capped = item.get('is_capped', False) 
        
        # 🟢 FIX 2: ใช้สถานะ is_passed ที่คำนวณจาก Engine (ซึ่งรวม Dependency Check แล้ว)
        final_is_passed = item.get('is_passed', item.get('is_pass', False)) 
        
        # ดึงเหตุผลดิบ
        reason_text_raw = item.get('reason_for_result', item.get('reason', 'ไม่ระบุเหตุผล'))
        
        # 2. กำหนดสถานะการแสดงผล
        if is_capped:
            status_text = "⚠️ FAIL (CAPPED)"
            status_color = COLOR_ACCENT
            capped_text = "YES"
            # แสดงเหตุผลเดิม แต่เพิ่มคำเตือน dependency
            reason_text = f"⚠️ [CAPPED] Dependency Failure (L{level-1} failed). (Raw Reason: {reason_text_raw})"
        elif final_is_passed:
            status_text = "✅ PASS"
            status_color = COLOR_GOOD
            capped_text = "-"
            reason_text = reason_text_raw
        else: # Final FAIL (Raw Score FAIL)
            status_text = "❌ FAIL"
            status_color = COLOR_BAD
            capped_text = "-"
            reason_text = reason_text_raw
        
        # 3. Policy exception for 1.1 L5 (V26 logic preserved)
        pdca_achieved_calc_sum = sum(v for k, v in pdca.items() if k in ['P', 'D', 'C', 'A'] and isinstance(v, int))
        
        # *** ตรงนี้มีการดึงค่า pdca_score_required ที่มาจาก seam_assessment.py แล้ว ***
        pdca_required = item.get('pdca_score_required', 8) 
        
        # Note: status_text needs to be '✅ PASS' to trigger this policy check
        if str(sub_id) == '1.1' and level == 5 and status_text == '✅ PASS': 
            if pdca_achieved_calc_sum == 4 and pdca_required == 4: 
                # ใช้เหตุผลดิบสำหรับ Policy Exception
                reason_text = f"**[POLICY EXCEPTION]** {reason_text_raw}" 

        # 4. PDCA Score (Achieved/Required) - V29 FIX PRESERVED
        # ดึงค่า 'score_achieved' จาก JSON โดยตรงตามที่ระบบประเมินได้คำนวณไว้ (CRITICAL FIX)
        score_achieved = item.get('score_achieved')
        
        if score_achieved is not None and isinstance(score_achieved, (int, float)):
            pdca_achieved_display = score_achieved
        else:
            # Fallback to calculation if 'score_achieved' is missing or invalid
            pdca_achieved_display = pdca_achieved_calc_sum
        
        # Format Achieved/Required Score text
        pdca_achieved_display = int(pdca_achieved_display)
        pdca_required = int(pdca_required) # ใช้ค่าที่ดึงมาจาก backend
        
        if pdca_achieved_display > 0 or pdca_required > 0:
            score_text = f"{pdca_achieved_display} / {pdca_required}" 
        else:
            score_text = "- / -"

        # --- Populate 9 Columns ---
        
        # Col 0: Level
        add_paragraph(cells[0], f"L{level}", size=TABLE_FONT_SIZE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Col 1: Status
        add_paragraph(cells[1], status_text, size=TABLE_FONT_SIZE, bold=True, color=status_color, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Col 2: PDCA Score (Achieved/Required)
        add_paragraph(cells[2], score_text, size=TABLE_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Col 3-6: P, D, C, A scores (New addition)
        add_paragraph(cells[3], str(pdca.get('P', '-')), size=TABLE_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(cells[4], str(pdca.get('D', '-')), size=TABLE_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(cells[5], str(pdca.get('C', '-')), size=TABLE_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(cells[6], str(pdca.get('A', '-')), size=TABLE_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Col 7: CAPPED
        capped_color_display = COLOR_ACCENT if is_capped else COLOR_NEUTRAL
        add_paragraph(cells[7], capped_text, size=TABLE_FONT_SIZE, bold=is_capped, color=capped_color_display, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Col 8: Reason/Analysis
        add_paragraph(cells[8], reason_text, size=TABLE_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph("")


def _add_pdca_legend(doc: Document):
    """เพิ่มคำอธิบายความหมายของคะแนน PDCA และสถานะต่างๆ ในตาราง"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    
    # 1. คำอธิบายคะแนน PDCA (0, 1, 2)
    r1 = p.add_run("คำอธิบายคะแนน PDCA: ")
    set_font_for_run(r1, size=12, bold=True, color=COLOR_HEADER)
    
    r2 = p.add_run("0 = ไม่พบหลักฐาน | 1 = หลักฐานเบื้องต้น/ไม่สมบูรณ์ (Initial) | 2 = หลักฐานสมบูรณ์/เป็นระบบ (Systematic)")
    set_font_for_run(r2, size=12, color=COLOR_NEUTRAL)

    # 2. คำอธิบายความหมายของ PDCA Cycle (P, D, C, A)
    p_pdca_def = doc.add_paragraph()
    p_pdca_def.paragraph_format.space_before = Pt(3)
    
    r_pdca_head = p_pdca_def.add_run("PDCA Cycle (ความหมาย): ")
    set_font_for_run(r_pdca_head, size=12, bold=True, color=COLOR_HEADER)
    
    r_p_def = p_pdca_def.add_run("P (Plan) - วางแผนและกำหนดกระบวนการ | ")
    set_font_for_run(r_p_def, size=12, color=COLOR_NEUTRAL)
    
    r_d_def = p_pdca_def.add_run("D (Do) - นำกระบวนการไปปฏิบัติจริง | ")
    set_font_for_run(r_d_def, size=12, color=COLOR_NEUTRAL)
    
    r_c_def = p_pdca_def.add_run("C (Check) - ตรวจสอบ/วัดผลการปฏิบัติงาน | ")
    set_font_for_run(r_c_def, size=12, color=COLOR_NEUTRAL)
    
    r_a_def = p_pdca_def.add_run("A (Act) - ปรับปรุง/แก้ไขตามผลการตรวจสอบ")
    set_font_for_run(r_a_def, size=12, color=COLOR_NEUTRAL)
    
    # 3. คำอธิบายสถานะ (Lvl Status)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(3)
    
    r3 = p2.add_run("Lvl Status: ")
    set_font_for_run(r3, size=12, bold=True, color=COLOR_HEADER)
    
    r4 = p2.add_run(
        "PASS/FAIL/CAPPED แสดงผลลัพธ์การประเมินต่อระดับ L1-L5. "
        "FAIL หมายถึง LLM ไม่พบหลักฐานเพียงพอ | PASS หมายถึง LLM พบหลักฐานเพียงพอ | "
    )
    set_font_for_run(r4, size=12, color=COLOR_NEUTRAL)

    # 4. คำอธิบาย CAPPED
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(3)
    
    r5 = p3.add_run("⚠️ CAPPED: ")
    set_font_for_run(r5, size=12, bold=True, color=COLOR_ACCENT)
    
    r6 = p3.add_run("แสดงว่าระดับนี้ได้รับ Raw PASS แต่ไม่สามารถนับเป็นระดับวุฒิภาวะเต็มได้ เนื่องจากมีระดับต่ำกว่า FAIL")
    set_font_for_run(r6, size=12, color=COLOR_NEUTRAL)

def _add_action_plan(doc: Document, sub_meta: Dict[str, Any]):
    """เพิ่มส่วนแผนปฏิบัติการ (Action Plan)"""
    action_plan = sub_meta.get('action_plan') or []
    
    if action_plan and any(plan.get('Actions') for plan in action_plan):
        add_paragraph(doc, "แผนปฏิบัติการและคำแนะนำการปรับปรุงหลักฐาน (Action Plan & Evidence Improvement):", bold=True, color=COLOR_HEADER, size=16)
        
        for plan in action_plan:
            goal = plan.get('Goal', 'ปรับปรุงเกณฑ์ย่อยนี้')
            p_goal = add_paragraph(doc, f"  • 🎯 เป้าหมายหลัก: {goal}", bold=True, size=14, color=COLOR_NEUTRAL)
            p_goal.paragraph_format.left_indent = Inches(0.2)
            
            actions = plan.get('Actions', [])
            
            if actions:
                for action in actions:
                    failed_lvl = action.get('Failed_Level', 'N/A')
                    target_type = action.get('Target_Evidence_Type', 'N/A')
                    rec = action.get('Recommendation', '(ไม่ระบุคำแนะนำ)')
                    
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    p.paragraph_format.space_before = Pt(6)
                    
                    r1 = p.add_run(f"    - Gap (L{failed_lvl} ไม่ผ่าน): ")
                    set_font_for_run(r1, size=14, bold=True, color=COLOR_BAD)
                    
                    r2 = p.add_run(f"ต้องการหลักฐานประเภท: ")
                    set_font_for_run(r2, size=14, bold=True)
                    
                    r3 = p.add_run(target_type)
                    set_font_for_run(r3, size=14, bold=True, color=COLOR_ACCENT)
                    
                    r4 = p.add_run("\n      📝 คำแนะนำเชิงรูปธรรม: ")
                    set_font_for_run(r4, size=14, bold=True, color=COLOR_HEADER)
                    
                    r5 = p.add_run(rec)
                    set_font_for_run(r5, size=14, color=COLOR_NEUTRAL)
                    
                    p.paragraph_format.left_indent = Inches(0.4)
            doc.add_paragraph("")
    else:
        add_paragraph(doc, "👍 เกณฑ์ย่อยนี้บรรลุเป้าหมายสูงสุด (L5) แล้ว หรือไม่มีข้อมูล Action Plan (กรุณาตรวจสอบ)", color=COLOR_GOOD, size=14)
    doc.add_paragraph("")

def _add_raw_evidence_table(doc: Document, stmts: List[Dict[str, Any]], sub_meta: Dict[str, Any]):
    """
    เพิ่มตารางแสดงหลักฐานอ้างอิงดิบ (Raw Evidence) 
    
    ปรับปรุง: ย้าย LLM Context Summary และ Source(s) ไปอยู่ในแถวที่รวมคอลัมน์
    เพื่อเพิ่มความสามารถในการอ่าน
    """
    add_paragraph(doc, "Raw Evidence / Statements (หลักฐานอ้างอิง + แหล่งที่มา)", bold=True, size=14)
    
    # 1. กำหนดจำนวนคอลัมน์หลักเป็น 4
    N_COLS = 4 
    ev_table = doc.add_table(rows=1, cols=N_COLS)
    ev_table.style = 'Table Grid'
    ev_table.allow_autofit = True
    
    # 2. กำหนดหัวข้อคอลัมน์และขนาด (4 คอลัมน์หลักเท่านั้น)
    # Headers เดิม: ["Statement (Level)", "Result", "Reason / Analysis", "LLM Context Summary", "Recommendation / Implication", "Source(s)"]
    # Headers ใหม่:
    headers = ["Statement (Level)", "Result", "Reason / Analysis", "Recommendation / Implication"]
    widths = [Inches(2.5), Inches(0.6), Inches(2.5), Inches(1.5)] # ปรับขนาดความกว้างตามจำนวนคอลัมน์ใหม่
    
    for i, h in enumerate(headers):
        cell = ev_table.rows[0].cells[i]
        add_paragraph(cell, h, size=TABLE_FONT_SIZE, bold=True)
        try:
            cell.width = widths[i]
        except:
             pass

    # 3. เตรียมข้อมูล Action Plan เพื่อดึงคำแนะนำสำหรับระดับที่ FAIL
    action_plan_list = sub_meta.get('action_plan', []) 
    action_plan_actions = []
    for plan in action_plan_list:
        if isinstance(plan, dict) and 'Actions' in plan:
            action_plan_actions.extend(plan['Actions'])
    
    highest_lvl = sub_meta.get('highest_full_level', 0)
    
    for rec in sorted(stmts, key=lambda r: r.get('level', 0)):
        
        # ------------------ แถวที่ 1: Main Statement Row (4 คอลัมน์) ------------------
        cells_main = ev_table.add_row().cells # มี 4 เซลล์ (Index 0-3)
        
        for cell in cells_main:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP 

        lvl = rec.get('level', '-')
        pass_flag = rec.get('is_passed', rec.get('is_pass', False)) 
        is_capped = rec.get('is_capped', False)
        
        # --- Column 1 (Index 1): Result ---
        if is_capped:
            pass_text = "⚠️ CAPPED"
            color_row = COLOR_ACCENT
        elif pass_flag:
            pass_text = "✅ PASS"
            color_row = COLOR_GOOD
        else:
            pass_text = "❌ FAIL"
            color_row = COLOR_BAD
            
        add_paragraph(cells_main[1], pass_text, size=TABLE_FONT_SIZE, bold=True, color=color_row, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # --- Column 0 (Statement) ---
        statement_text = rec.get('statement', '(no statement)')
        text0 = f"L{lvl}: {statement_text[:SNIPPET_MAX_CHARS]}{'...' if len(statement_text) > SNIPPET_MAX_CHARS else ''}"
        add_paragraph(cells_main[0], text0, size=14) 
        
        # --- Column 2 (Reason / Analysis) ---
        reason_text = rec.get('reason_for_result', rec.get('reason', '-')) 
        
        # Policy exception logic (คงไว้ซึ่ง Logic เดิม)
        sub_id = rec.get('sub_criteria_id')
        pdca = rec.get('pdca_breakdown', {})
        pdca_achieved_calc_sum = sum(v for k, v in pdca.items() if k in ['P', 'D', 'C', 'A'] and isinstance(v, int))
        
        is_policy_pass_exception = (
            pass_flag and 
            str(sub_id) == '1.1' and 
            str(lvl) == '5' and 
            all(v == 1 for k, v in pdca.items() if k in ['P', 'D', 'C', 'A'] and isinstance(v, int)) and
            pdca_achieved_calc_sum == 4
        )
        
        if is_policy_pass_exception:
            reason_text = (
                f"🚨 [POLICY EXCEPTION] เกณฑ์ 1.1 (นโยบาย) นี้ถือว่า **PASS L5** แม้คะแนน PDCA เป็น 1,1,1,1 (รวม 4 คะแนน) เนื่องจากหลักฐานที่พบครอบคลุมหัวข้อเชิงนโยบายระดับสูงสุดได้อย่างชัดเจน "
                f"ซึ่งถือว่าเพียงพอต่อการยืนยันการมีอยู่ของนโยบาย/ทิศทางระดับ 5 ตามความคาดหวังของเกณฑ์นี้"
            )

        if is_capped:
            reason_text = f"⚠️ [LLM PASS / SEQUENCING CAPPED] {reason_text} (ระดับต่ำกว่าไม่ผ่าน จึงไม่นับเป็นระดับวุฒิภาวะเต็ม)"
            
        add_paragraph(cells_main[2], reason_text[:SNIPPET_MAX_CHARS], size=14) 

        # --- Column 3 (Recommendation / Implication) ---
        if pass_flag and not is_capped: 
            rec_text = f"รักษาสถานะ/ดำเนินการสู่ L{int(lvl)+1}"
            rec_color = COLOR_GOOD
        elif is_capped:
            rec_text = "⚠️ ต้องแก้ไขหลักฐานระดับต่ำกว่า (L3 หรือ L4) ก่อนจึงจะนับระดับนี้"
            rec_color = COLOR_ACCENT
        else:
            failed_lvl = rec.get('level')
            specific_action = next((a for a in action_plan_actions if str(a.get('Failed_Level', '')) == str(failed_lvl)), None)
            
            if specific_action:
                rec_text = f"สร้างหลักฐาน: {specific_action.get('Target_Evidence_Type')}\nแนะนำ: {specific_action.get('Recommendation')}"
            else:
                rec_text = f"แก้ไขหลักฐาน L{lvl} ให้ผ่านเกณฑ์ (ดูรายละเอียดใน Action Plan)"
            rec_color = COLOR_BAD
        
        add_paragraph(cells_main[3], rec_text, size=14, color=rec_color) 

        # ------------------ แถวที่ 2: Merged Detail Row (การแก้ไขหลัก) ------------------
        
        cells_detail = ev_table.add_row().cells
        
        # 1. รวมคอลัมน์ทั้งหมด (Merge all N_COLS cells: 0 ถึง 3)
        merged_cell = cells_detail[0].merge(cells_detail[N_COLS - 1])
        
        # 2. เตรียมข้อมูล LLM Context Summary
        llm_summary_text = (
            rec.get('llm_summarized_context') or
            safe_get(rec, 'llm_summary_full_result', 'summary') or
            '(ไม่พบ LLM Summary)'
        )
        
        # 3. เตรียมข้อมูล Sources
        sources = extract_sources_from_stmt(rec)
        source_list = '\n• ' + '\n• '.join(sources) if sources else '(ไม่มีแหล่งที่มา)'

        # 4. ใส่ข้อมูลรวม (Summary + Sources) ลงใน Merged Cell โดยใช้ Run เพื่อทำ Bold
        p = merged_cell.add_paragraph()
        
        # 4.1 Header: LLM Context Summary
        run = p.add_run("LLM Context Summary:")
        run.bold = True
        
        # 4.2 Content: LLM Summary Text
        p.add_run(f"\n{llm_summary_text}\n\n")

        # 4.3 Header: Sources
        run = p.add_run("Sources:")
        run.bold = True
        
        # 4.4 Content: Source List
        p.add_run(f"\n{source_list}")
        
    doc.add_paragraph("")

def build_detailed_section(doc: Document, raw_llm_results: List[Dict[str, Any]], sub_results: List[Dict[str, Any]], target_level_value: int):
    """สร้างส่วนผลการวิเคราะห์อย่างละเอียด (Part B) แบ่งตามเกณฑ์ย่อย"""
    add_heading_center(doc, "Part B: Detailed Findings (PDCA per Level) & Evidence", level=1)
    
    sub_map = {s.get('sub_criteria_id'): s for s in sub_results}
    
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for stmt in raw_llm_results:
        sid = stmt.get('sub_criteria_id', 'UNKNOWN')
        grouped.setdefault(sid, []).append(stmt)

    for sub_id in sorted(grouped.keys()):
        stmts = grouped[sub_id]
        sub_meta = sub_map.get(sub_id, {})
        
        # 1. เพิ่มหัวข้อและ Meta Data
        _add_sub_criteria_meta(doc, sub_meta.get('sub_criteria_id') or sub_id, sub_meta, target_level_value)
        
        # 2. ตารางสรุปสถานะรายระดับ (V30: รวม PDCA Breakdown เข้าไปแล้ว)
        _add_level_status_summary_table(doc, sub_meta) 
        
        # 3. คำอธิบายตาราง (Legend)
        _add_pdca_legend(doc) 
        doc.add_paragraph("")
        
        # 4. Action Plan 
        _add_action_plan(doc, sub_meta)
        
        # 5. Raw Evidence 
        _add_raw_evidence_table(doc, stmts, sub_meta) 

        doc.add_page_break()

# -------------------------
# MAIN BUILDER / ฟังก์ชันหลัก (v31)
# -------------------------
def build_report(results_file: str, sub_filter: Optional[str], enabler: Optional[str], out_path: Optional[str]):
    """
    ฟังก์ชันหลักในการอ่านข้อมูลและสร้างไฟล์รายงาน DOCX
    """
    data = load_json(results_file)
    if not data:
         print(f"❌ Error: ไม่สามารถโหลดไฟล์ JSON '{results_file}' ได้ หรือไฟล์ว่างเปล่า")
         return
         
    summary = data.get('summary', {})
    sub_results = data.get('sub_criteria_results', [])
    
    sub_id = sub_filter or safe_get(summary, 'sub_criteria_id')
    
    raw_llm_results = []
    for sr in sub_results:
        # V27: ใช้ raw_results_ref เพื่อให้ได้ข้อมูลครบถ้วนสำหรับทุกระดับ
        raw_llm_results.extend(sr.get('raw_results_ref', [])) 
    
    # กรองข้อมูลตาม Sub-criteria ที่ระบุ (ถ้ามี)
    if sub_filter:
        sub_results_filtered = [s for s in sub_results if str(s.get('sub_criteria_id', '')).upper() == sub_filter.upper()]
        
        if len(sub_results_filtered) == 1:
            sub_results = sub_results_filtered
        
        # ปรับ raw_llm_results ให้ตรงกับ sub_criteria ที่เลือก
        raw_llm_results = [r for r in raw_llm_results if str(r.get('sub_criteria_id', '')).upper() == sub_filter.upper()]
            
        if not sub_results and not raw_llm_results:
             print(f"❌ ไม่พบข้อมูลสำหรับเกณฑ์ย่อย '{sub_filter}' ในไฟล์ JSON")
             return
    
    if not raw_llm_results:
        # Fallback to older structure if raw_results_ref is missing (legacy support)
        raw_llm_results = data.get('raw_llm_results', [])

    if not raw_llm_results:
        print(f"❌ ไม่พบข้อมูล raw_results_ref หรือ raw_llm_results ที่ถูกต้องในไฟล์ JSON '{results_file}'")
        return

    # กำหนดชื่อ Enabler ฉบับเต็ม
    enabler_id = (enabler or summary.get('enabler') or summary.get('enabler_id') or "GENERIC").upper()
    enabler_full = SEAM_ENABLER_MAP.get(enabler_id, f"Enabler {enabler_id}")

    # กำหนดชื่อไฟล์เอาต์พุต
    if out_path:
        out_file = out_path
    else:
        safe_sub = sub_filter.replace('.', '_') if sub_filter else "ALL"
        out_file = os.path.join(EXPORT_DIR, f"{enabler_id}_Comprehensive_Report_{safe_sub}_{REPORT_DATE}.docx")

    ensure_output_dir(out_file)

    doc = Document()
    
    try:
        style = doc.styles['Normal']
        style.font.name = THAI_FONT
        style.element.rPr.rFonts.set(qn('w:eastAsia'), THAI_FONT)
        style.font.size = Pt(DEFAULT_FONT_SIZE)
    except:
        pass

    target_lvl_val = safe_get(summary, 'target_level', 0)

    # สร้าง Part A: สรุปภาพรวม
    build_overall_summary(doc, summary, sub_results, enabler_full)
    
    # สร้าง Part B: รายละเอียดและหลักฐานอ้างอิง
    build_detailed_section(doc, raw_llm_results, sub_results, target_lvl_val)

    doc.save(out_file)
    print(f"✅ สร้างรายงานสำเร็จ: {out_file}")

def main():
    """ฟังก์ชันเริ่มต้นเมื่อรันสคริปต์"""
    parser = argparse.ArgumentParser(description="Generate SEAM assessment report (DOCX) with readable PDCA & Raw Evidence")
    parser.add_argument("results_file", help="Path to JSON results file")
    parser.add_argument("--sub", help="Sub-criteria id to filter (e.g., 2.2)", default=None)
    parser.add_argument("--enabler", help="Enabler id (e.g., KM)", default=None)
    parser.add_argument("--output", help="Output file path (optional)", default=None)
    args = parser.parse_args()
    
    # ตั้งค่า default (หากไม่ได้ใส่ filter แต่ชื่อไฟล์บ่งชี้)
    if not args.sub and '1.1' in args.results_file:
         args.sub = '1.1'
    elif not args.sub and '3.2' in args.results_file:
         args.sub = '3.2'
         
    if not args.enabler and ('KM' in args.results_file or '2.2' in args.results_file):
         args.enabler = 'KM'

    build_report(args.results_file, args.sub, args.enabler, args.output)

if __name__ == "__main__":
    # ฟังก์ชันหลักสำหรับการรันสคริปต์
    main()