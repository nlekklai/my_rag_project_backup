#report_generator.py
import os
import argparse
import json
from datetime import datetime
import re # <<< เพิ่มการ Import re สำหรับฟังก์ชัน clean_for_display
from typing import Dict, Any, Optional, List

# ==================== Essential Imports for DOCX ====================
# ***ต้องติดตั้ง python-docx ก่อน: pip install python-docx***
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT # For Landscape
from collections import defaultdict
from docx.shared import Inches, Pt, RGBColor, Cm # <<< Cm ต้องถูก Import


# ==================== Global Constants ====================
REPORT_DATE = datetime.now().strftime("%Y-%m-%d")
EXPORT_DIR = "reports" # Default output directory

# NOTE: Since 'from config.global_vars import SEAM_ENABLER_MAP' is used in the original, 
# but the variable is not defined in this file, a fallback is provided here 
# to make the script runnable without the external config file.
try:
    from config.global_vars import SEAM_ENABLER_MAP
except ImportError:
    # Fallback definition if SEAM_ENABLER_MAP is not available externally
    SEAM_ENABLER_MAP = {
        "CG": "1 การกำกับดูแลที่ดีและการนำองค์กร (Corporate Governance & Leadership)",
        "SP": "2 การวางแผนเชิงยุทธศาสตร์ (Strategic Planning)",
        "RM&IC": "3 การบริหารความเสี่ยงและการควบคุมภายใน (Risk Management & Internal Control)",
        "SM": "4.1 การมุ่งเน้นผู้มีส่วนได้ส่วนเสีย (Stakeholder Management)",
        "CM": "4.2 การมุ่งเน้นลูกค้า (Customer Management)",
        "DT": "5 การพัฒนาเทคโนโลยีดิจิทัล (Digital Technology)",
        "HCM": "6 การบริหารทุนมนุษย์ (Human Capital Management)",
        "KM": "7.1 การจัดการความรู้ (Knowledge Management)",
        "IM": "7.2 นวัตกรรม (Innovation Management)",
        "IA": "8 การตรวจสอบภายใน (Internal Audit)"
    }

THAI_FONT_NAME = "Angsana New" # Standard Thai Font for DOCX

# =========================================================================
# 1. Utility Functions
# =========================================================================

def setup_output_folder(output_path_or_dir):
    """จัดการสร้างโฟลเดอร์ Output"""
    if os.path.isdir(output_path_or_dir):
        output_dir = output_path_or_dir
    else:
        output_dir = os.path.dirname(output_path_or_dir)
        if not output_dir:
            output_dir = EXPORT_DIR
    
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"📦 สร้างโฟลเดอร์: {output_dir}")

def load_data(file_path, data_name):
    """โหลดไฟล์ JSON"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            print(f"✅ โหลดไฟล์ {data_name} สำเร็จ: {file_path}")
            return json.load(f)
    except FileNotFoundError:
        print(f"❌ ไม่พบไฟล์: {file_path}")
        return None
    except json.JSONDecodeError:
        print(f"❌ ไฟล์ {file_path} มีปัญหาในการถอดรหัส JSON")
        return None

def flatten_raw_data(raw_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """แปลง Raw Data ให้อยู่ในรูปแบบ List ของ Statements"""
    if isinstance(raw_data, list):
        return raw_data
    return []

def set_cell_color(cell, color_rgb):
    """Utility function to set cell background color (Hex format: 0xRRGGBB)."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), f'{color_rgb:06x}')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_thai_font(run, size_pt=Pt(11)):
    """Utility function to apply Thai font (Angsana New) and size to a run."""
    run.font.name = THAI_FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), THAI_FONT_NAME)
    run.font.size = size_pt

def set_landscape(doc):
    """
    ตั้งค่าหน้ากระดาษเป็นแนวนอน (Landscape) ปรับขนาด A4 และขอบกระดาษให้เล็กลง
    
    Returns:
        float: ความกว้างรวมของตารางที่เหมาะสมสำหรับหน้านี้ (Inches)
    """
    section = doc.sections[-1]
    
    # 1. ตั้งค่าแนวนอน
    section.orientation = WD_ORIENT.LANDSCAPE
    
    # 2. ปรับขนาดหน้ากระดาษเป็น A4 Landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    # 3. ตั้งค่าขอบกระดาษใหม่สำหรับแนวนอน (0.5 นิ้ว ทุกด้าน)
    margin = Inches(0.5) 
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    
    # 4. ความกว้างที่เหลือสำหรับตาราง (ประมาณ 11.69 - 0.5 - 0.5 = 10.69 นิ้ว)
    return Inches(10.5) # กำหนดความกว้างตารางที่ปลอดภัย

# <<< START: ฟังก์ชัน clean_for_display ที่เพิ่มเข้ามา >>>
def clean_for_display(retrieved_text: str) -> str:
    """
    Cleans up the segmented text retrieved from the vector store for final display 
    by removing '|' and cleaning up excessive spaces/artifacts.
    """
    if not retrieved_text:
        return ""
        
    # 1. แทนที่ตัวแบ่งคำ '|' ด้วยช่องว่างปกติ
    text = retrieved_text.replace('|', ' ') 

    # 2. ลบช่องว่างที่เกิดจากการทำความสะอาดมากเกินไป และรวมช่องว่างที่เกินสองช่องให้เหลือช่องว่างเดียว
    text = re.sub(r'\s{2,}', ' ', text)
    
    # 3. ลบอักขระขึ้นบรรทัดใหม่ที่เกินความจำเป็น (เหลือไม่เกิน 2 บรรทัดติดกัน)
    text = re.sub(r'(\n|\r\n|\r){2,}', '\n\n', text)
    
    # 4. ลบช่องว่างก่อน/หลังเครื่องหมายวรรคตอนที่อาจหลงเหลืออยู่จาก Segmentation/OCR
    text = re.sub(r'\s*([.,:;])\s*', r'\1 ', text) # จัดการกับ . , : ; 
    text = text.replace(' )', ')').replace('( ', '(')
    
    return text.strip()
# <<< END: ฟังก์ชัน clean_for_display ที่เพิ่มเข้ามา >>>

# =========================================================================
# 2. DOCX Formatting and Setup Functions
# =========================================================================

def setup_document(doc):
    """กำหนดรูปแบบเอกสาร สร้างหน้าปก และเพิ่มสารบัญ placeholder"""
    
    # 1. สร้างหน้าปก (Title Page)
    doc.add_paragraph() # เพิ่มย่อหน้าว่างเพื่อช่วยจัดตำแหน่ง
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("รายงานผลการประเมินเชิงกลยุทธ์\n(Strategic Assessment Report)")
    title_run.font.size = Pt(28)
    title_run.bold = True
    set_thai_font(title_run, Pt(28)) # Apply font to title

    # Subtitle
    subtitle = doc.add_paragraph("\nSE-AM AI Assessment Project")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"\nประเมิน ณ วันที่: {REPORT_DATE}")
    subtitle_run.font.size = Pt(18)
    set_thai_font(subtitle_run, Pt(18)) # Apply font to subtitle
    
    # Page Break หลังหน้าปก
    doc.add_page_break()
    
    # 2. สารบัญ (Table of Contents) - ใช้ Field Code เพื่อให้อัปเดตอัตโนมัติใน MS Word
    
    # Ensure font is applied to the heading title
    heading = doc.add_heading("สารบัญ (Table of Contents)", level=1)
    for run in heading.runs: 
        set_thai_font(run, Pt(16))
    
    # Field code for TOC: { TOC \o "1-3" \h \z \u }
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    set_thai_font(run, Pt(11)) 
    
    # Insert TOC field code manually (required for python-docx)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u' # Field Instruction
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    # Page Break หลังสารบัญ
    doc.add_page_break()

# =========================================================================
# 3. DOCX Content Generation Functions
# =========================================================================

def generate_overall_summary_docx(doc, summary_data, enabler_name_full):
    # SECTION 1
    heading = doc.add_heading(f"[SECTION 1] สรุปผลการประเมิน {enabler_name_full} โดยรวม", level=1)
    for run in heading.runs: set_thai_font(run, Pt(16))
    
    overall = summary_data.get("Overall", {})
    
    # Data Extraction
    enabler_id = overall.get("enabler", "N/A")
    score = overall.get("total_weighted_score", 0.0)
    possible_score = overall.get("total_possible_weight", 0.0)
    progress_percent = overall.get("overall_progress_percent", 0.0)
    maturity_score = overall.get("overall_maturity_score", 0.0)
    
    # Create Summary Table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = "ตัวขับเคลื่อน (Enabler):"
    table.cell(0, 1).text = f"{enabler_id} ({enabler_name_full})"
    table.cell(1, 0).text = "คะแนนรวมถ่วงน้ำหนักที่ได้:"
    table.cell(1, 1).text = f"{score:.2f} / {possible_score:.2f}"
    table.cell(2, 0).text = "เปอร์เซ็นต์ความคืบหน้าโดยรวม:"
    
    progress_cell = table.cell(2, 1)
    progress_cell_p = progress_cell.paragraphs[0]
    progress_cell_run = progress_cell_p.add_run(f"{progress_percent:.2f}%")
    progress_cell_run.bold = True
    
    # Set color for progress
    if progress_percent >= 70:
        color = RGBColor(0x00, 0x80, 0x00) # Green
    elif progress_percent >= 50:
        color = RGBColor(0xFF, 0x8C, 0x00) # DarkOrange
    else:
        color = RGBColor(0xFF, 0x00, 0x00) # Red
    progress_cell_run.font.color.rgb = color
        
    table.cell(3, 0).text = "คะแนนวุฒิภาวะโดยรวม (Maturity Score):"
    table.cell(3, 1).text = f"{maturity_score:.2f}"
    
    # Apply Thai Font to all table cells
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_thai_font(run)

    doc.add_paragraph("\n")


def generate_executive_summary_docx(doc, summary_data):
    # SECTION 2
    heading = doc.add_heading("[SECTION 2] สรุปสำหรับผู้บริหาร (Executive Summary)", level=1)
    for run in heading.runs: set_thai_font(run, Pt(16))
    
    overall = summary_data.get("Overall", {})
    p1 = doc.add_paragraph(f"✅ คะแนนรวม: {overall.get('total_weighted_score', 0.0):.2f} / {overall.get('total_possible_weight', 0.0):.2f}")
    p2 = doc.add_paragraph(f"✅ ร้อยละความสำเร็จ: {overall.get('overall_progress_percent', 0.0):.2f}%")
    p3 = doc.add_paragraph(f"✅ ระดับความเป็นผู้ใหญ่: {overall.get('overall_maturity_score', 0.0):.2f}")
    for p in [p1, p2, p3]:
         for run in p.runs: set_thai_font(run)

    
    # --- 2.1 จุดแข็งที่โดดเด่น (Top Strengths) ---
    heading2_1 = doc.add_heading("📈 จุดแข็งที่โดดเด่น (Top Strengths):", level=2)
    for run in heading2_1.runs: set_thai_font(run, Pt(14))
    
    top_strengths = sorted(
        summary_data.get("SubCriteria_Breakdown", {}).items(), 
        key=lambda item: item[1].get('score', 0), 
        reverse=True
    )[:3]

    if top_strengths:
        for sub_id, data in top_strengths:
            p = doc.add_paragraph(f"• [{sub_id}] {data.get('topic')} (คะแนน: {data.get('score', 0):.2f}/{data.get('weight', 0):.2f})", style='List Bullet')
            for run in p.runs: set_thai_font(run)
    else:
        p = doc.add_paragraph("ไม่พบจุดแข็งที่ชัดเจน.")
        for run in p.runs: set_thai_font(run)
    
    # --- 2.2 จุดที่ควรพัฒนา (Development Areas) ---
    heading2_2 = doc.add_heading("🚨 จุดที่ควรพัฒนา (Development Areas):", level=2)
    for run in heading2_2.runs: set_thai_font(run, Pt(14))
    
    development_areas = []
    for sub_id, data in summary_data.get("SubCriteria_Breakdown", {}).items():
        if data.get("development_gap") and data.get('highest_full_level', 0) < 5:
            gap_text = (
                f"• [{sub_id}] {data.get('topic')} "
                f"(ระดับสูงสุดผ่าน: L{data.get('highest_full_level', 0)})"
            )
            development_areas.append(gap_text)
            
    if development_areas:
        for item in development_areas:
            p = doc.add_paragraph(item)
            p.style = 'List Bullet'
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            for run in p.runs: set_thai_font(run)
    else:
        p = doc.add_paragraph("✅ ไม่พบ Gap ที่สำคัญในผลการประเมินนี้.")
        for run in p.runs: set_thai_font(run)
    
    doc.add_paragraph("\n")


def generate_sub_criteria_status_docx(doc, summary_data):
    """
    SECTION 3: Sub-Criteria Status & Gap (สร้างตารางละเอียดราย Level พร้อม Color-coding)
    *** เพิ่มคอลัมน์สุดท้าย: คะแนนรวม (Score/Weight) ***
    """
    heading = doc.add_heading("[SECTION 3] สถานะการประเมินรายเกณฑ์ย่อยและ Gap", level=1)
    for run in heading.runs: set_thai_font(run, Pt(16))
    
    p = doc.add_paragraph("ตารางนี้แสดงระดับความสำเร็จ (Pass Ratio) ในแต่ละ Level เพื่อระบุ Gap ที่ชัดเจน")
    for run in p.runs: set_thai_font(run)

    breakdown = summary_data.get("SubCriteria_Breakdown", {})
    if not breakdown:
        doc.add_paragraph("ไม่พบข้อมูลเกณฑ์ย่อยสำหรับการประเมิน.")
        return []

    # 1. Setup Table
    # *** 8 คอลัมน์ ***
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    
    # Set column widths (Portrait Mode)
    from docx.shared import Inches
    table.columns[0].width = Inches(0.5)  # ID
    table.columns[1].width = Inches(2.5)  # Topic
    for i in range(5):
        table.columns[i+2].width = Inches(0.8) # L1-L5 (รวม 4.0 นิ้ว)
    # *** คอลัมน์ใหม่: คะแนนรวม ***
    table.columns[7].width = Inches(1.0)  # Score/Weight (รวม 8.0 นิ้ว)

    # Headers
    hdr = table.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "เกณฑ์ย่อย (Pass Level)"
    for i in range(5):
        hdr[i+2].text = f"L{i+1} Ratio"
    # *** เพิ่ม Header คอลัมน์สุดท้าย ***
    hdr[7].text = "คะแนนรวม (Score/Weight)" 
    
    # Apply shading and font to header
    for cell in hdr:
         set_cell_color(cell, 0xC0C0C0) # Light Grey
         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
         for run in cell.paragraphs[0].runs: set_thai_font(run, Pt(10)) 

    gap_criteria_docx = []

    # 2. Populate Rows
    for sub_id, data in breakdown.items():
        if not data.get("development_gap", False) and data.get("highest_full_level", 0) == 5:
            continue # Skip fully passed criteria

        row = table.add_row().cells
        
        # Column 1: ID
        row[0].text = sub_id
        
        # Column 2: Topic & Highest Full Level
        topic_text = f"{data.get('topic')}\n(Highest Full: L{data.get('highest_full_level', 0)})"
        row[1].text = topic_text
        
        if data.get("development_gap", False):
            gap_criteria_docx.append(sub_id)
            
        pass_ratios = data.get("pass_ratios", {})
        
        # Columns 3-7: L1-L5 Pass Ratio
        for i in range(5):
            level = str(i + 1)
            ratio = pass_ratios.get(level, 0.0)
            ratio_percent = f"{ratio*100:.0f}%"

            cell = row[i + 2]
            cell.text = ratio_percent
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Color Coding
            if ratio >= 1.0:
                set_cell_color(cell, 0xD9EAD3) # Light Green (Full Pass)
            elif ratio > 0.0:
                set_cell_color(cell, 0xFFE5CC) # Light Orange (Partial Pass)
            else: # ratio == 0.0
                set_cell_color(cell, 0xF4CCCC) # Light Red (Zero Pass)

        # *** คอลัมน์ 8: คะแนนรวม (Score/Weight) ***
        score = data.get('score', 0.0)
        weight = data.get('weight', 0.0)
        score_text = f"{score:.2f}/{weight:.2f}"
        
        score_cell = row[7]
        score_cell.text = score_text
        score_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Apply font to all cells in the row and highlight failed ID
        for cell in row:
            for run in cell.paragraphs[0].runs: 
                set_thai_font(run, Pt(10))
        
        # Highlight Gap for Action Plan
        if data.get("development_gap", False):
            row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Red ID

    doc.add_paragraph("\n")
    return gap_criteria_docx


def generate_action_plan_report_docx(doc, final_summary_data, gap_criteria_docx):
    # SECTION 4
    heading = doc.add_heading("[SECTION 4] รายงานรายละเอียดแผนปฏิบัติการเพื่อปิดช่องว่าง (Action Plan)", level=1)
    for run in heading.runs: set_thai_font(run, Pt(16))
    
    if not gap_criteria_docx:
         p = doc.add_paragraph("ไม่มีเกณฑ์ย่อยที่ต้องสร้างแผนปฏิบัติการ (Action Plan)")
         for run in p.runs: set_thai_font(run)
         return
         
    # Iterate through SubCriteria that require action plan
    for sub_id in gap_criteria_docx:
        sub_data = final_summary_data.get("SubCriteria_Breakdown", {}).get(sub_id, {})
        action_plans = final_summary_data.get("Action_Plans", {}).get(sub_id, [])
        
        # Sub-heading for the Sub-Criteria
        heading2 = doc.add_heading(f"• เกณฑ์ย่อย {sub_id}: {sub_data.get('topic')} (Highest Full Level: L{sub_data.get('highest_full_level', 0)})", level=2)
        for run in heading2.runs: set_thai_font(run, Pt(14))
        
        # L5 Summary (Insight)
        p1 = doc.add_paragraph(f"💡 ข้อมูลเชิงลึกหลักฐานระดับ L5 (เป้าหมายสูงสุด):")
        for run in p1.runs: set_thai_font(run)
        p2 = doc.add_paragraph(f"   - สรุปหลักฐาน: {sub_data.get('evidence_summary_L5', {}).get('summary', 'ไม่พบข้อมูลสรุปหลักฐาน L5')}")
        for run in p2.runs: set_thai_font(run)
        
        # Action Plan Table
        if action_plans:
            p = doc.add_paragraph(f"แผนปฏิบัติการ ({len(action_plans)} Phase):")
            for run in p.runs: set_thai_font(run)
            
            for phase_data in action_plans:
                p_phase = doc.add_paragraph(f"🛠️ เฟส/ขั้นตอน: {phase_data.get('Phase')}")
                p_goal = doc.add_paragraph(f"🎯 เป้าหมายหลัก: {phase_data.get('Goal')}")
                for p in [p_phase, p_goal]:
                    for run in p.runs: set_thai_font(run)
                
                # Table for Actions
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = "คำแนะนำ (Recommendation)"
                hdr[1].text = "หลักฐานเป้าหมาย (Evidence Type)"
                hdr[2].text = "ตัวชี้วัดสำคัญ (Key Metric)"
                
                for cell in hdr:
                    set_cell_color(cell, 0xC0C0C0)
                    for run in cell.paragraphs[0].runs: set_thai_font(run, Pt(10))

                for action in phase_data.get("Actions", []):
                    row = table.add_row().cells
                    row[0].text = action.get("Recommendation", "")
                    row[1].text = action.get("Target_Evidence_Type", "")
                    row[2].text = action.get("Key_Metric", "")
                    for cell in row:
                        for run in cell.paragraphs[0].runs: set_thai_font(run, Pt(10))
        else:
             p = doc.add_paragraph("✅ ไม่พบ Action Plan ที่ระบุไว้ในไฟล์ Summary Data.")
             for run in p.runs: set_thai_font(run)
             
        doc.add_paragraph("\n")



def generate_raw_details_report_docx(doc, final_raw_data, table_width=Inches(7.2)):
    """
    SECTION 5: Raw Details (สำหรับ RawDetails.docx)
    *** การแก้ไข: กำหนดความกว้างคอลัมน์ Snippet เป็น 6 cm และเพิ่ม Sub-Header row ***
    """
    
    if not final_raw_data or len(final_raw_data) == 0:
        p = doc.add_paragraph("⚠️ ไม่พบหลักฐานเชิงลึก (Raw Details) สำหรับเกณฑ์ย่อยนี้ หรือข้อมูลไม่พร้อมสำหรับการประเมิน.")
        p.style = 'Intense Quote'
        for run in p.runs: set_thai_font(run)
        return

    # 1. Setup Table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.autofit = True # ใช้ Autofit เพื่อให้คอลัมน์อื่นปรับอัตโนมัติ
    
    # 🟢 [FIX 3] กำหนดความกว้างคอลัมน์ Snippet (Index 5) เป็น 6 cm
    try:
        # กำหนดความกว้างของคอลัมน์สุดท้าย (Index 5: หลักฐาน/บริบท) เป็น 6 cm
        table.columns[5].width = Cm(6)
    except NameError:
        # Fallback หาก Cm ไม่ได้ถูก Import หรือ Error
        # 6 cm ≈ 2.362 inches
        table.columns[5].width = Inches(2.362)
    except Exception as e:
        # จัดการข้อผิดพลาดอื่นๆ ที่อาจเกิดขึ้นในการกำหนดความกว้างคอลัมน์
        print(f"⚠️ Warning: Failed to set specific column width: {e}") 
        pass
        
    # Headers 
    headers = ["เกณฑ์ย่อย | Statement ID", "ผล", "Statement / Standard", "เหตุผล/วิเคราะห์", "แหล่งที่มา", "หลักฐาน/บริบท (Snippet)"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_color(cell, 0xC0C0C0) # Light Grey
        for run in cell.paragraphs[0].runs:
            run.bold = True
            set_thai_font(run, Pt(8)) # ใช้ Pt(8) สำหรับ Header 

    # 2. Populate Rows (New structure with Sub-Headers)
    
    # Group data by sub_criteria_id
    grouped_data = defaultdict(list)
    for item in final_raw_data:
        sub_id = item.get('sub_criteria_id', 'N/A')
        grouped_data[sub_id].append(item)
    
    # Iterate through grouped data to insert sub-headers
    for sub_id, statements in grouped_data.items():
        if not statements:
            continue
            
        sub_criteria_name = statements[0].get('sub_criteria_name', 'ชื่อเกณฑ์ย่อยไม่ระบุ')
        
        # --- Insert Sub-Header Row ---
        header_row = table.add_row().cells
        
        # Merge cells from 0 to 5 (all 6 columns)
        merged_cell = header_row[0].merge(header_row[5])
        
        # Content
        merged_cell.text = f"⚙️ เกณฑ์ย่อย {sub_id}: {sub_criteria_name}"
        
        # Formatting
        set_cell_color(merged_cell, 0xEAF1DD) # A light green/grey for grouping
        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = merged_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs: 
            set_thai_font(run, Pt(9))
            run.bold = True
            
        # --- Insert regular statement rows for this sub_criteria ---
        for item in statements:
            row = table.add_row().cells

            # Status
            is_passed = item.get("pass_status", False)
            
            # Content for Col 1 (Status) - ใช้ข้อความเต็ม
            status_th = item.get('status_th')
            status_text = status_th or ("ผ่าน" if is_passed else "ไม่ผ่าน")
            status_text_full = f"✅ {status_text}" if is_passed else f"❌ {status_text}"
            
            # Determine Cell Color for FAIL (Red)
            fill_color = 0xFFFFFF # White
            if not is_passed:
                fill_color = 0xF4CCCC # Light Red
            
            # Combine Statement and Standard
            statement_full = f"S: {item.get('statement', '')}\nStd: {item.get('standard', '')}"
            
            # Combine Source names
            source_names_list = [src.get('source_name', '') for src in item.get('retrieved_sources_list', [])]
            source_names = "\n".join(source_names_list[:2]) # แสดงแค่ 2 แหล่งแรก
            if len(source_names_list) > 2:
                source_names += f"\n(...{len(source_names_list)-2} more)"
                
            # Combined ID: e.g., "1.1 | L1-S1"
            statement_id_short = f"L{item.get('level', '')}-S{item.get('statement_number', '')}"
            combined_id = f"{sub_id} | {statement_id_short}"
            
            # Row content
            content = [
                combined_id,
                status_text_full,
                statement_full,
                item.get('reason', ''),
                source_names,
                clean_for_display(item.get('context_retrieved_snippet', '')) 
            ]

            for i, text in enumerate(content):
                cell = row[i]
                cell.text = str(text)
                set_cell_color(cell, fill_color)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                
                # Apply Thai Font to all text in cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_thai_font(run, Pt(8)) # ใช้ Pt(8) สำหรับ Body 
                        if i == 1: # Status column
                            run.bold = True
                            if not is_passed:
                                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Red text for FAIL

    doc.add_paragraph("\n")


# ==========================
# 4. MAIN EXECUTION
# ==========================
def main():
    """ฟังก์ชันหลักในการสร้างรายงานทั้งหมด"""
    
    parser = argparse.ArgumentParser(description="Generate Comprehensive Assessment Reports.")
    parser.add_argument("--mode", choices=["all", "sub"], default="all", help="all: Generate full report. sub: Generate report for a specific sub-criteria.")
    parser.add_argument("--sub", type=str, help="SubCriteria ID (e.g., 2.2) if mode=sub.")
    parser.add_argument("--summary_file", type=str, required=True, help="Path to the Strategic/Summary JSON file.")
    parser.add_argument("--raw_file", type=str, required=True, help="Path to the Raw Details JSON file.")
    parser.add_argument("--output_path", type=str, default=os.path.join(EXPORT_DIR, "KM_Comprehensive_Report"), help="Output directory and base filename prefix (e.g., reports/KM_Comprehensive_Report).")
    
    args = parser.parse_args()
    
    # 1. จัดการ Folder Output และแยก Directory
    output_dir = os.path.dirname(args.output_path)
    if not output_dir:
         output_dir = EXPORT_DIR
    setup_output_folder(output_dir)
    
    # 2. โหลดไฟล์
    summary_data = load_data(args.summary_file, "Strategic/Summary Data")
    raw_data = load_data(args.raw_file, "Raw Details Data") 
    
    if not summary_data:
        print("🚨 ไม่สามารถสร้างรายงานได้เนื่องจากไฟล์ Summary Core Data ไม่พร้อม")
        return
    
    # --- 3. ดึง ENABLER และกำหนดค่าเริ่มต้น ---
    enabler_id = summary_data.get("Overall", {}).get("enabler", "KM").upper() 
    enabler_name_full = SEAM_ENABLER_MAP.get(enabler_id, f"Unknown Enabler ({enabler_id})")
    
    final_summary_data = summary_data
    final_raw_data = raw_data
    
    # --- 4. การจัดการชื่อไฟล์และการกรองข้อมูล ---
    
    # 4.1. กำหนด Base Filename Prefix
    base_prefix = os.path.basename(args.output_path)
    if not base_prefix or base_prefix == "KM_Comprehensive_Report":
        base_prefix = f"{enabler_id}_Comprehensive_Report"
        
    # 4.2. การกรองข้อมูลสำหรับโหมด 'sub'
    if args.mode == "sub" and args.sub:
        sub_id = args.sub.upper()
        print(f"🔹 โหมด: รายงานเฉพาะเกณฑ์ย่อย {sub_id} สำหรับ {enabler_name_full}")
        
        # กรอง Summary Data
        if sub_id in summary_data.get("SubCriteria_Breakdown", {}):
            final_summary_data = {
                "Overall": summary_data.get("Overall",{}),
                "SubCriteria_Breakdown": {sub_id: summary_data["SubCriteria_Breakdown"][sub_id]},
                "Action_Plans": {sub_id: summary_data.get("Action_Plans",{}).get(sub_id,[])}
            }
        else:
             print(f"⚠️ ไม่พบข้อมูลสำหรับเกณฑ์ย่อย {sub_id} ใน Summary Data. ใช้ข้อมูลทั้งหมดของ Summary แทน.")
            
        # กรอง Raw Data 
        if raw_data is not None:
            all_statements = flatten_raw_data(raw_data)
            # Filter by sub_criteria_id (assuming it's present in each statement object)
            filtered_statements = [
                stmt for stmt in all_statements 
                if stmt.get("sub_criteria_id", "").upper() == sub_id
            ]
            final_raw_data = filtered_statements if filtered_statements else None
            
        # อัปเดต Base Prefix สำหรับโหมด Sub
        base_prefix = f"{enabler_id}_Report_{sub_id}"
    
    else:
        print(f"🔹 โหมด: รายงานฉบับเต็มสำหรับ {enabler_name_full}")
        
    # 4.3. กำหนดชื่อ Output สุดท้าย (รวมวันที่)
    final_base_name = f"{base_prefix}_{REPORT_DATE}"
    
    # สร้างชื่อไฟล์แยกตามประเภทและรวมวันที่
    strategic_path = os.path.join(output_dir, f"{final_base_name}_Strategic.docx")
    detail_path = os.path.join(output_dir, f"{final_base_name}_RawDetails.docx")
    # output_txt_path = os.path.join(output_dir, f"{final_base_name}.txt") # Omitted for brevity

    # --- A. การสร้างไฟล์ DOCX (แยกเป็น 2 ไฟล์: Strategic และ Raw Details) ---
    
    # 1. สร้าง Strategic Report (Sections 1-4)
    print(f"\nกำลังสร้างไฟล์ DOCX [Strategic Report]...")
    strategic_doc = Document()
    setup_document(strategic_doc) # สร้างหน้าปกและสารบัญ
    
    # SECTION 1: Overall Summary
    generate_overall_summary_docx(strategic_doc, final_summary_data, enabler_name_full) 
    # SECTION 2: Executive Summary
    generate_executive_summary_docx(strategic_doc, final_summary_data)
    # SECTION 3: Sub-Criteria Status & Gap (ถูกแก้ไขแล้ว)
    gap_criteria_docx = generate_sub_criteria_status_docx(strategic_doc, final_summary_data)
    # SECTION 4: Action Plan Report
    generate_action_plan_report_docx(strategic_doc, final_summary_data, gap_criteria_docx)

    # บันทึกไฟล์ Strategic Report
    strategic_doc.save(strategic_path)
    print(f"🎉 สร้างไฟล์ DOCX [Strategic Report] สำเร็จ! บันทึกที่: {strategic_path}")


    # 2. สร้าง Raw Details Working Document (Section 5)
    print(f"กำลังสร้างไฟล์ DOCX [Raw Details]...")
    detail_doc = Document()
    
    # *** ตั้งค่าหน้ากระดาษเป็นแนวนอนและปรับความกว้างตาราง ***
    landscape_table_width = set_landscape(detail_doc) 
    
    # Heading เฉพาะสำหรับ Raw Details
    detail_doc.add_heading(f"[SECTION 5] รายงานหลักฐานเชิงลึก (Raw Details) - {enabler_name_full} ({REPORT_DATE})", level=1)
    
    # SECTION 5: Raw Details
    # ใช้ landscape_table_width เป็นค่าที่ตั้งไว้ (10.5 นิ้ว) 
    generate_raw_details_report_docx(detail_doc, final_raw_data, table_width=landscape_table_width) 

    # บันทึกไฟล์ Raw Details
    detail_doc.save(detail_path)
    print(f"🎉 สร้างไฟล์ DOCX [Raw Details] สำเร็จ! บันทึกที่: {detail_path}")

if __name__ == "__main__":
    main()